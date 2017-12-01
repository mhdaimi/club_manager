#!/usr/bin/python
# -*- coding: utf-8 -*-
import os,csv,time
import Tkinter as tk
from tkinter import ttk
from ttk import Progressbar
import tkFileDialog, tkMessageBox
import sqlite3, string, re, datetime
from pyPdf import PdfFileWriter, PdfFileReader
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
from reportlab.lib.styles import getSampleStyleSheet,ParagraphStyle, TA_CENTER
from reportlab.lib.units import inch,mm
from reportlab.lib.pagesizes import letter
#import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document, text
from docx.shared import Pt
from docx.enum.text import *

background_color='PeachPuff3'
display_bg="light grey"
label_bg = "grey"
global all_member_data, all_finance_update_recs, all_finance_history_recs, all_comm_recs, all_comm_history_recs
all_finance_update_recs=[]
all_member_data = []
all_finance_history_recs = []
all_comm_recs=[]
all_comm_history_recs=[]

if os.environ['PATH'].find("iCLS Client")>=0:
    os.environ['PATH'] = "".join([it for it in os.environ['PATH'].split(";") if not it.find("iCLS Client")>0])

class ScrollableFrame(tk.Frame):
    def __init__(self, master, **kwargs):
        tk.Frame.__init__(self, master, **kwargs)
        self.vscrollbar = tk.Scrollbar(self, orient=tk.VERTICAL)
        self.vscrollbar.pack(side='right', fill="y",  expand="false")
        self.canvas = tk.Canvas(self,
                                bg=display_bg, bd=1,
                                height=300,
                                width=1290,
                                highlightthickness=1,
                                yscrollcommand=self.vscrollbar.set)
        self.canvas.pack(side="left", fill="both", expand="true")
        self.vscrollbar.config(command=self.canvas.yview)
        self.canvas.xview_moveto(0)
        self.canvas.yview_moveto(0)
        self.interior = tk.Frame(self.canvas, **kwargs)
        self.canvas.create_window(0, 0, window=self.interior, anchor="nw")
        self.bind('<Configure>', self.set_scrollregion)
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        
    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(-1*(event.delta/120), "units")
        
    def set_scrollregion(self, event=None):
        """ Set the scroll region on the canvas"""
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))

class VereinManager(tk.Frame):
    def __init__(self, master):
        self.master = master
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=5, column=0)
        self.file_name = "verein_manager.db"
        self.dir_path = ''
        self.fields_list = ["Ref. No : ", "Name : ", "Street : ", "No : ", "PostCode : ", "City : ", "Phone No : ", "Date of Birth (dd/mm/yyyy) : ",
                            "Monthly Contribution : ", "Email : ", "Member Since (dd/mm/yyyy) : ", " Select Gender : ", "Comments : " ]
        self.conn=""
        self.check_config_file()
        
    def check_config_file(self):
        if os.path.isfile(os.path.join(os.path.expanduser("~"), 'verein_manager.config')):
            with open (os.path.join(os.path.expanduser("~"), 'verein_manager.config'), 'r') as fh:
                dir_path = fh.readlines()
                self.dir_path = dir_path[0]
                self.start_page()
        else:
            tkMessageBox.showinfo("Select Default Directory", "Please select a default directory to store your database file! ")
            dir_name = tkFileDialog.askdirectory(parent=self.master,title='Please select a directory')
            if dir_name:
                with open (os.path.join(os.path.expanduser("~"), 'verein_manager.config'), 'w') as fh:
                    fh.writelines(dir_name)
                    self.dir_path = dir_name
                    self.start_page()
            else:
                tkMessageBox.showerror("Directory Not Set", "Restart the program and select a directory to continue!")
                self.master.destroy()
    
#     def login_page(self):
#         self.first_frame.destroy()
#         self.first_frame = tk.Frame(self.master, bg=background_color)
#         self.first_frame.grid(row=5, column=0, padx=150, pady=75)
#         name_label = tk.Label(self.first_frame,bg=background_color,text="User Name :")
#         name_label.grid(row=0,column=0,sticky=tk.W,pady=10,columnspan=3)
#         self.user_name = tk.Entry(self.first_frame,width=17)
#         self.user_name.grid(row=0,column=1,sticky=tk.W,padx=1, columnspan=3)
#         pass_label = tk.Label(self.first_frame,bg=background_color,text="Password :")
#         pass_label.grid(row=1,column=0,sticky=tk.W,columnspan=3)
#         self.pass_word = tk.Entry(self.first_frame,width=17,show="*")
#         self.pass_word.grid(row=1,column=1,sticky=tk.W,padx=1, columnspan=3)
#         submit_but = tk.Button(self.first_frame, text="Submit", command=self.validate_user,width=12,height=1)
#         submit_but.grid(row=3,column=0,pady=25)
#         submit_but = tk.Button(self.first_frame, text="Exit", command=self.exit_prog,width=12,height=1)
#         submit_but.grid(row=3,column=1,padx=15)
    
    def validate_user(self):
        if self.user_name.get() == "":
            tkMessageBox.showwarning("Input Required", "User name cannot be blank")
        elif self.user_name.get() != "admin":
            tkMessageBox.showwarning("Invalid User", "Invalid User name")
        elif self.pass_word.get() == "":
            tkMessageBox.showwarning("Input Required", "Password cannot be blank")
        elif self.pass_word.get() != "Allah786!":
            tkMessageBox.showwarning("Incorrect Password", "Password is incorrect")
        else:
            self.start_page()
    
    def check_database(self):
        self.conn=sqlite3.connect(self.dir_path + '/' + self.file_name)
        self.cur = self.conn.cursor()
        self.cur.execute("""
                        CREATE table IF NOT EXISTS member_data(Id INT PRIMARY KEY, ref_no TEXT, name TEXT, street TEXT, house_no TEXT, post_code INT, city BLOB, phone TEXT,
                        date_of_birth TEXT, donation REAL, email TEXT, date_joined TEXT, comment TEXT, gender TEXT)
            """)
        
        self.conn.commit()
        self.cur.execute("""
                        CREATE table IF NOT EXISTS finance(Id INT, month TEXT, year TEXT, actual_amt REAL, donated_amt REAL, name TEXT, ref_no TEXT)
            """)
        self.conn.commit()
        
        self.cur.execute("""
                        CREATE table IF NOT EXISTS deleted_member_data(Id, ref_no TEXT, name TEXT, street TEXT, house_no INT, post_code INT, city BLOB, phone TEXT,
                        date_of_birth TEXT, donation REAL, email TEXT, date_joined TEXT, comment TEXT, gender TEXT, deleted TEXT)
            """)
        
        self.conn.commit()
        
        self.cur.execute("""
                        CREATE table IF NOT EXISTS comm_history(Id INT , ref_no TEXT, name TEXT, date_sent TEXT, message_sent BLOB, how_sent TEXT)
            """)
        
        self.conn.commit()
    
    def start_page(self,check=False):
        global all_member_data
        all_member_data=[]
        self.check_database()
        if check:
            self.checkbox_pane.destroy()
        self.first_frame.destroy()
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=5, column=0, padx=150, pady=75)
        b3 = tk.Button(self.first_frame,text="Members", command=self.members_page, height = 1, width = 15)
        b3.grid(row=8,column=1,padx=25, pady=40)
        b5 = tk.Button(self.first_frame,text="Finance", command=self.manage_finance, height = 1, width = 15)
        b5.grid(row=8,column=3, padx=25, pady=20)
        b7 = tk.Button(self.first_frame, text="Generate Report", command=self.generate_financial_report, height=1, width=15)
        b7.grid(row=8, column=4, padx=5, pady=10)
        b5 = tk.Button(self.first_frame,text="History", command=self.view_all_history, height = 1, width = 15)
        b5.grid(row=9,column=3, padx=25, pady=20)
        b6 = tk.Button(self.first_frame,text="Import/Export Data", command=self.import_export_data, height = 1, width = 15)
        b6.grid(row=9,column=1, padx=25, pady=20)
        b4 = tk.Button(self.first_frame,text="Quit", command=self.exit_prog, height = 1, width = 15)
        b4.grid(row=9,column=4, padx=25, pady=40)
        l1 = tk.Label(self.first_frame,text="Your database file is stored at: %s" %(self.dir_path + '/' + self.file_name))
        l1.grid(row=21, columnspan=5, padx=5, pady=180)
        
    def members_page(self):
        global all_member_data
        all_member_data=[]
        self.first_frame.destroy()
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=5, column=0, padx=10, pady=5)
        label_fr = tk.Frame(self.first_frame, bg=background_color)
        label_fr.grid(row=0, column=0, padx=10,sticky=tk.N+tk.W)
        self.recs_label=tk.Label(label_fr,bg=background_color,justify=tk.LEFT)
        self.recs_label.grid(row=0,column=0,columnspan=2, sticky=tk.N+tk.W, pady=5)
        search_frame = tk.Frame(self.first_frame, bg=background_color)
        search_frame.grid(row=0, column=2, padx=50,rowspan=5,sticky=tk.N+tk.W)
        self.create_search_by_field(search_frame)
        self.search_all_footer = tk.Label(search_frame, bg=background_color)
        self.search_all_footer.grid(row=0, column=1, padx=2, columnspan=2, pady=1)
        nd_label_fr = tk.Frame(self.first_frame, bg=background_color)
        nd_label_fr.grid(row=0, column=1, padx=75,sticky=tk.N+tk.W)
        self.no_data_label=tk.Label(nd_label_fr,bg=background_color,justify=tk.LEFT)
        self.no_data_label.grid(row=0,column=0,columnspan=4, pady=5)
        prev_nxt_fr = tk.Frame(self.first_frame, bg=background_color)
        prev_nxt_fr.grid(row=1, column=0, padx=10, pady=1,sticky=tk.N+tk.W)
        self.next_but = tk.Button(prev_nxt_fr, text = 'Next', height = 1, width = 12)
        self.prev_but = tk.Button(prev_nxt_fr, text = 'Previous', height = 1, width = 12)
        self.next_but.grid(row=1, column=0, padx=5, pady=1)
        self.prev_but.grid(row=1, column=1, padx=5, pady=1)
        self.next_but.config(state=tk.DISABLED)
        self.prev_but.config(state=tk.DISABLED)
        but_fr = tk.Frame(self.first_frame, bg=background_color)
        but_fr.grid(row=1, column=1, padx=75, rowspan=5,sticky=tk.N+tk.W)
        self.add_new_but = tk.Button(but_fr,text="Add New", command=self.add_member_data, height = 1, width = 12)
        self.add_new_but.grid(row=1,column=0,padx=5, pady=1)
        self.del_but = tk.Button(but_fr, text="Delete", command=self.delete_recs, height = 1, width = 12)
        self.del_but.grid(row=1, column=2, pady=1, padx=5)
        self.edit_but = tk.Button(but_fr, text="Edit", command=self.edit_record, height = 1, width = 12)
        self.edit_but.grid(row=1, column=1, pady=1,padx=5)
        self.show_but = tk.Button(but_fr,text="Show All Data", command=lambda: self.display_all_records(check_des=True), height = 1, width = 12)
        self.show_but.grid(row=2,column=0, padx=5, pady=5)
        self.prepare_but = tk.Button(but_fr, text = 'Prepare Letter', command =lambda: self.select_language(all_member_data,from_mem_page=True), height = 1, width = 12)
        self.prepare_but.grid(row=2, column=1, padx=5, pady=5)
        b4 = tk.Button(but_fr,text="Back", command=lambda: self.start_page(check=True), height = 1, width = 12)
        b4.grid(row=2,column=2, padx=5, pady=5)
        self.display_all_records()
        
    def import_export_data(self):
        self.first_frame.destroy()
        self.first_frame = tk.Frame(self.master,bg=background_color)
        self.first_frame.grid(row=0)
        but_frame = tk.Frame(self.first_frame,bg=background_color)
        but_frame.grid(row=1, column=1,padx=100,pady=100)
        prog_frame = tk.Frame(self.first_frame,bg=background_color)
        prog_frame.grid(row=1, column=2, sticky=tk.W)
        self.progress = Progressbar(prog_frame, orient="vertical", length=250,
                                        mode="determinate")
        
        but = tk.Button(but_frame,text="Import CSV", command=self.select_import_file, height=1,width=15)
        but.grid(row=5,column=1, pady=5,padx=10)
        b6 = tk.Button(but_frame,text="Export CSV", command=self.export_data, height = 1, width = 15)
        b6.grid(row=5,column=2, padx=10, pady=20)
        backbut = tk.Button(but_frame,text="Back", command=self.start_page, height=1,width=15)
        backbut.grid(row=6,column=1,sticky=tk.W,padx=10)
        
    def add_member_data(self):
        self.checkbox_pane.destroy()
        self.first_frame.destroy()
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=1, column=1, padx=200, pady=25)
        self.create_input_fields(field_list=self.fields_list)
        button_1 = tk.Button(self.first_frame, text="Submit", command=self.add_record, height = 1, width = 15)
        button_1.grid(row=15, column=2, padx=5, pady=25)
        button_2 = tk.Button(self.first_frame, text="Reset", command=self.reset_fields, height = 1, width = 15)
        button_2.grid(row=15, column=3, padx=5, columnspan=2)
        button_3 = tk.Button(self.first_frame, text="Back", command=self.members_page, height = 1, width = 15)
        button_3.grid(row=16, column=2, padx=5, pady=5, sticky=tk.N)
    
    def add_record(self,from_import=""):
        self.all_fields_verified = False
        if not from_import:
            self.check_empty_fields()
        if self.no_empty_field:
            if not from_import: 
                self.validate_input()
            else:
                error_msg = self.validate_input(from_import=from_import)
                if error_msg:
                    return error_msg
        if self.all_fields_verified and self.no_empty_field:
            if self.donation.get() == "":
                donation = 0
            else:
                donation = float(self.donation.get())
            self.cur.execute("SELECT * FROM member_data where Id = (SELECT MAX(Id) from member_data)")
            rec = self.cur.fetchone() 
            if rec:
                max_id = rec[0]
                id = max_id + 1
            else:
                id =1
                
            sql = "INSERT into member_data VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)"
            parameters = [id, self.ref_no.get().zfill(4), string.capwords(self.f_name.get()), string.capwords(self.street.get()), self.house_no.get(),  
                          int(self.post_code.get()), string.capwords(self.city.get()), self.phone.get(), self.date_of_birth.get(), donation, self.email.get(),
                           self.date_joined.get(), self.comment.get("1.0",tk.END).strip("\n"), string.capwords(self.selected_gender.get())]
            try:
                self.cur.execute(sql, parameters)
                self.conn.commit()
            except Exception, msg:
                tkMessageBox.showerror("Error", "%s occurred while adding new record! " %str(msg))
            else:
                if not from_import:
                    tkMessageBox.showinfo("Add Data", "Record added successfully")
                else:
                    return "Record added successfully\n"
                self.clear_input_fields()    
    
    def reset_fields(self):
        self.first_frame.destroy()
        self.add_member_data()    
    
    def display_all_records(self, check_des=False):
        self.disable_red_border_clear_text("All")
        self.search_all_footer.config(text="")
        self.search_field_footer.config(text="")
        if check_des:
            self.checkbox_pane.destroy()
        self.checkbox_pane = ScrollableFrame(self.master, bg=display_bg, relief=tk.RIDGE)
        self.checkbox_pane.grid(row=0, column=0, rowspan=5, columnspan=5, padx=5, pady=5)
        self.cur.execute(""" SELECT * from member_data""")
        all_data = self.cur.fetchall()
        self.initialise_fields(display=True)
        self.start_num=0
        self.data_length=25
        self.max_num=25
        self.max_len = len(all_data)
        if all_data:
            self.no_data_label.config(text="")
            if self.max_len <= self.max_num:
                self.max_num = self.max_len
                self.next_but.config(state=tk.DISABLED)
                self.prev_but.config(state=tk.DISABLED)
            for i in range(self.start_num,self.max_num,1):
                self.display_recs(i,all_data[i])
            self.recs_label.config( text="Displaying Records : %d to %d of %d" %(self.start_num+1,i+1,self.max_len))
            if len(all_data) > self.data_length:
                self.next_but.config(state=tk.NORMAL)
                self.next_but.config(command=lambda: self.show_next(all_data, display=True))
                self.prev_but.config(command =lambda: self.show_previous(all_data, display=True))
        else:
            self.display_recs("","")
            self.recs_label.config(text="")
            self.del_but.config(state=tk.DISABLED)
            self.edit_but.config(state=tk.DISABLED)
            self.prepare_but.config(state=tk.DISABLED)
            self.show_but.config(state=tk.DISABLED)
            self.no_data_label.config(text="Database has no records to display!") 
    
    def create_search_by_field(self, search_frame):
        self.s_name = tk.Entry(search_frame, width=15)
        self.s_city = tk.Entry(search_frame, width=15)
        self.s_phone = tk.Entry(search_frame, width=15)
        self.s_ref_no= tk.Entry(search_frame, width=15)
        search_fields = ["Name", "City", "Phone", "Ref. No"]
        self.var = tk.StringVar()
        rows=0
        self.radio_butt = []
        for i, field in zip(range(len(search_fields)),search_fields):
            rows += 1
            self.radio_butt.append(tk.Radiobutton(search_frame, text=field, variable=self.var, value=field, height = 1, width = 13, 
                                               command=self.selected_rb, bg=background_color ))
            self.radio_butt[i].grid(row=rows, pady=3, column=1,sticky=tk.E)
        self.radio_butt[0].select()
        self.s_name.grid(row=1, column=2, padx=2)
        self.s_city.grid(row=2, column=2)
        self.s_phone.grid(row=3, column=2)
        self.s_ref_no.grid(row=4, column=2)
        self.s_name.config(state="normal")
        self.s_phone.config(state="disabled", disabledbackground=label_bg)
        self.s_city.config(state="disabled", disabledbackground=label_bg)
        self.s_ref_no.config(state="disabled", disabledbackground=label_bg)
        self.label_name = tk.Label(search_frame, text="", width=20, bg=background_color, justify=tk.LEFT)
        self.label_name.grid(row=1, column=3, sticky=tk.E)
        self.label_city = tk.Label(search_frame, text="", width=20, bg=background_color, justify=tk.LEFT)
        self.label_city.grid(row=2, column=3, sticky=tk.E)
        self.label_phone = tk.Label(search_frame, text="", width=20, bg=background_color, justify=tk.LEFT)
        self.label_phone.grid(row=3, column=3, sticky=tk.E)
        self.label_ref_no = tk.Label(search_frame, text="", width=20, bg=background_color, justify=tk.LEFT)
        self.label_ref_no.grid(row=4, column=3, sticky=tk.E)
        button_5 = tk.Button(search_frame, text="Search By ", command=lambda: self.search_by_field(self.var.get()), height = 1, width = 12)
        button_5.grid(row=1, column=0, pady=1, padx=5)
        button_6 = tk.Button(search_frame, text="Clear Field ", command=lambda: self.disable_red_border_clear_text("All"), height = 1, width = 12)
        button_6.grid(row=3, column=0, pady=1, padx=5)
        self.search_field_footer = tk.Label(search_frame, bg=background_color, justify=tk.LEFT)
        self.search_field_footer.grid(row=0, column=0, pady=10, columnspan=4, stick=tk.W)
        
    def search_by_field(self, field_name):
        field_validate = False
        self.search_field_footer.config(text="")
        self.search_all_footer.config(text="")
        self.no_data_label.config(text="")
        if field_name == "Name":
            self.disable_red_border_clear_text(field_name)
            if self.s_name.index("end") == 0:
                self.label_name.config(text="<-- Please enter Name", fg="red")
                self.s_name.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            else:
                self.s_name.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                field_validate = True
                self.label_name.config(text="")
        elif field_name == "City" :
            self.disable_red_border_clear_text(field_name)
            if self.s_city.index("end") == 0:
                self.label_city.config(text="<-- Please enter City", fg="red")
                self.s_city.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            else:
                self.s_city.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                field_validate = True
                self.label_city.config(text="")
        elif field_name == "Phone":
            self.disable_red_border_clear_text(field_name)
            if self.s_phone.index("end") == 0:
                self.s_phone.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                self.label_phone.config(text="<-- Please enter Phone", fg="red")
            else:
                self.s_phone.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                field_validate = True
                self.label_phone.config(text="")
        elif field_name == "Ref. No":
            self.disable_red_border_clear_text(field_name)
            if self.s_ref_no.index("end") == 0:
                self.s_ref_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                self.label_ref_no.config(text="<-- Please enter Ref. No", fg="red")
            else:
                self.s_ref_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                field_validate = True
                self.label_ref_no.config(text="")
        if field_validate:
            self.fetch_record_and_display()
    
    def delete_recs(self):
        global all_member_data
        if not all_member_data:
            tkMessageBox.showinfo('Delete', "Please select record(s) to delete!")
        else:
            var = tkMessageBox.askyesno('Delete', "Are you sure you want to delete %s record(s)?" %(str(len(all_member_data))))
            if var:
                for each_rec in all_member_data:
                    values = each_rec.split(";")
                    id = values[0]
                    try:
                        self.cur.execute("DELETE from member_data where Id = ?",[(id)])
                        self.conn.commit()
                    except Exception, msg:
                        tkMessageBox.showerror("Error", "%s occurred while deleting record! " %str(msg))
                    else:
                        street,h_no = values[3].split("\n")
                        post_c,city = values[4].split("\n")
                        sql = "INSERT into deleted_member_data VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)"
                        parameters = [values[0],values[1],values[2],street,h_no,post_c,city,values[5],values[8],values[6],values[7],values[9],
                                      "Deleted on: %s"%str("{:%d.%m.%Y}".format(datetime.datetime.now())),values[11],"True"]
                        try:
                            self.cur.execute(sql,parameters)
                            self.conn.commit()
                        except Exception, msg:
                            tkMessageBox.showerror("Error", "%s occurred while deleting record! " %str(msg))
                tkMessageBox.showinfo('Delete', "Successfully deleted %s record(s)" %(str(len(all_member_data))))
                self.checkbox_pane.destroy()
                self.display_all_records()
            else:
                tkMessageBox.showinfo("Delete", "Member deletion cancelled!")

    def edit_record(self):
        global all_member_data
        if not all_member_data:
            tkMessageBox.showinfo('Edit', "Please select a record to edit!")
        elif len(all_member_data) > 1:
            tkMessageBox.showwarning('Edit', "Please select only one record to edit!")
        else:
            self.checkbox_pane.destroy()
            self.first_frame.destroy()
            each_val = all_member_data[0].split(";")
            self.first_frame = tk.Frame(self.master, bg=background_color)
            self.first_frame.grid(row=1, column=1, padx=200, pady=25, sticky=tk.E)
            self.create_input_fields(field_list=self.fields_list, from_edit=True)
            button_1 = tk.Button(self.first_frame, text="Save", command=lambda: self.save_record(int(each_val[0])), height = 1, width = 15)
            button_1.grid(row=15, column=2, padx=5, pady=50)
            button_3 = tk.Button(self.first_frame, text="Back", command=self.members_page, height = 1, width = 15)
            button_3.grid(row=15, column=3, padx=5, pady=25, columnspan=5)
            self.ref_no.insert(0,each_val[1].zfill(4))
            self.ref_no.config(state=tk.DISABLED)
            self.f_name.insert(0,each_val[2])
            street,h_no = each_val[3].split("\n")
            self.street.insert(0,street)
            self.house_no.insert(0,h_no)
            post_code, city = each_val[4].split("\n")
            self.post_code.insert(0,post_code)
            self.city.insert(0,city)            
            self.phone.insert(0,each_val[5])
            self.date_of_birth.insert(0,each_val[8])
            self.email.insert(0,each_val[7])
            self.donation.insert(0,each_val[6])
            self.date_joined.insert(0,each_val[9])
            self.comment.insert(tk.END,each_val[10])
            self.gender_entry.insert(0, each_val[11])
            self.gender_entry.config(state=tk.DISABLED)

    def manage_finance(self, kill_checkbox=""):
        global all_member_data
        all_member_data=[]
        self.first_frame.destroy()
        if kill_checkbox:
            self.checkbox_pane.destroy()
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=5, column=0)
        but_frame = tk.Frame(self.first_frame, bg=background_color)
        but_frame.grid(row=0, column=0, padx=250, pady=150)
        b1 = tk.Button(but_frame, text="Update", command=self.update_finance_data, height=1, width=15)
        b1.grid(row=1, column=1, padx=25, pady=25)
        b4 = tk.Button(but_frame, text="Check Payments", command=self.check_current_payments, height=1, width=15)
        b4.grid(row=1, column=2, padx=5, pady=10)
        b3 = tk.Button(but_frame, text="Back", command=self.start_page, height=1, width=15)
        b3.grid(row=2, column=1, padx=5, pady=10)
    
    def comm_with_member(self,check_des=""):
        if check_des:
            self.checkbox_pane.destroy()
        self.first_frame.destroy()  
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=5, column=0)
        field_fr = tk.Frame(self.first_frame,bg=background_color)
        field_fr.grid(row=0,column=0, padx=150,pady=75,sticky=tk.S)
        but_fr = tk.Frame(self.first_frame,bg=background_color)
        but_fr.grid(row=1,column=0, padx=5,pady=10,sticky=tk.N)
        label = tk.Label(field_fr,text="Reference No :", bg=background_color)
        label.grid(row=0,column=0)
        self.comm_ref_no_by_user = tk.Entry(field_fr,width=12)
        self.comm_ref_no_by_user.grid(row=0,column=1)
        view_but = tk.Button(but_fr, text = 'View History', command = self.view_comm_history, height = 1, width = 12)
        view_but.grid(row=0, column=1, padx=5,sticky=tk.N)
        select_none_but = tk.Button(but_fr, text = 'Back', command =self.view_all_history, height = 1, width = 12)
        select_none_but.grid(row=0, column=2, padx=5,sticky=tk.N)
        
    def display_comm_data(self,ix,text):    
        label0 = tk.Label(self.checkbox_pane.interior,width=5, borderwidth=1, relief=tk.SUNKEN)
        label2 = tk.Label(self.checkbox_pane.interior, text="Name", width=25,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label3 = tk.Label(self.checkbox_pane.interior, text="Reference No.", width=13,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label4 = tk.Label(self.checkbox_pane.interior, text="Deleted", width=13,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label0.grid(row=0, column=0, sticky=tk.E)
        label2.grid(row=0, column=1, sticky=tk.E)
        label3.grid(row=0, column=2, sticky=tk.E)
        label4.grid(row=0, column=3, sticky=tk.E)
        if text:
            id=text[0]
            ref_no = text[1]
            name = text[2]
            self.comm_hist_cb.append(tk.Radiobutton(self.checkbox_pane.interior, variable=self.comm_hist_cb_v, value=name+";"+str(id),
                                          command=self.comm_hist_cb_checked, bg=display_bg))
            self.comm_hist_cb[ix].grid(row=ix+1, column=0, sticky=tk.W, padx=5, pady=5)
            if len(text) == 15:
                label3=tk.Label(self.checkbox_pane.interior, text="Y", bg=display_bg)
                label3.grid(row=ix+1, column=3)
            label1=tk.Label(self.checkbox_pane.interior, text="%s" %str(name), bg=display_bg, wraplength=125)
            label2=tk.Label(self.checkbox_pane.interior, text="%s" %str(ref_no).zfill(4), bg=display_bg)
            label1.grid(row=ix+1, column=1)
            label2.grid(row=ix+1, column=2)
    
    def comm_hist_cb_checked(self):
        global all_comm_history_recs
        all_comm_history_recs=[]
        all_comm_history_recs.append(self.comm_hist_cb_v.get())
        
    def view_comm_history(self):
        if self.comm_ref_no_by_user.get() == "":
            tkMessageBox.showwarning("View History", "Reference number cannot be blank!")
        elif not self.comm_ref_no_by_user.get().isdigit() :
            tkMessageBox.showwarning("View History", "Please enter only numbers [0-9]")
        elif len(self.comm_ref_no_by_user.get()) > 4 :
            tkMessageBox.showwarning("View History", "Length of Reference number cannot be greater than 4")
        else:
            stmt = "Select * from comm_history WHERE ref_no=?"
            parameter =  [str(self.comm_ref_no_by_user.get()).zfill(4)]
            self.cur.execute(stmt,parameter)
            all_data = self.cur.fetchall()
            if all_data:
                ref_no = all_data[0][1]
                name = all_data[0][2]
                self.first_frame.destroy()
                self.first_frame = tk.Frame(self.master, bg=background_color)
                self.first_frame.grid(row=5, column=0)
                self.checkbox_pane = ScrollableFrame(self.master, bg=display_bg, relief=tk.RIDGE)
                self.checkbox_pane.grid(row=0, column=0, rowspan=5, columnspan=5, padx=5, pady=5)
                back_but = tk.Button(self.first_frame, text="Back", command=lambda: self.comm_with_member(check_des=True), width=15, height=1)
                back_but.grid(row=1, column=1)
                label1 = tk.Label(self.checkbox_pane.interior, text="Name", width=15,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                label2 = tk.Label(self.checkbox_pane.interior, text="Date Contacted", width=15,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                label3 = tk.Label(self.checkbox_pane.interior, text="Contacted Via", width=20,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                label4 = tk.Label(self.checkbox_pane.interior, text="Message", width=115,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                label1.grid(row=0, column=0, sticky=tk.E, pady=15)
                label2.grid(row=1, column=0, sticky=tk.E)
                label3.grid(row=1, column=1, sticky=tk.E)
                label4.grid(row=1, column=2, columnspan=10)
                for ix, each_rec in enumerate(all_data):
                    label1 = tk.Label(self.checkbox_pane.interior, text=each_rec[2], width=20,wraplength=125,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                    label2 = tk.Label(self.checkbox_pane.interior, text=each_rec[3], width=15,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                    label3 = tk.Label(self.checkbox_pane.interior, text=each_rec[5], width=20,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                    text = tk.Text(self.checkbox_pane.interior, width=100, height=2, wrap=tk.WORD )
                    label1.grid(row=0, column=1, sticky=tk.E)
                    label2.grid(row=ix+2, column=0, sticky=tk.E)
                    label3.grid(row=ix+2, column=1, sticky=tk.E)
                    text.grid(row=ix+2, column=2, columnspan=10)
                    text.insert(tk.END, each_rec[4])
            else:
                tkMessageBox.showinfo("View History", "No History found for member with reference no. %s " %str(self.comm_ref_no_by_user.get()).zfill(4))
        
    def check_current_payments(self,kill_checkbox=""):
        if kill_checkbox:
            self.checkbox_pane.destroy()
        self.first_frame.destroy()
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=5, column=0)
        but_frame = tk.Frame(self.first_frame, bg=background_color)
        but_frame.grid(row=1,column=1, padx=250, pady=150)
        select_all_but = tk.Button(but_frame, text = 'Defaulter', command =lambda: self.get_payment_data(defaulter=True,save_msg=True), height = 1, width = 15)
        select_all_but.grid(row=1, column=1, padx=25, pady=25)
        select_none_but = tk.Button(but_frame, text = 'Non Defaulter', command =lambda: self.get_payment_data(defaulter=False), height = 1, width = 15)
        select_none_but.grid(row=1, column=2, padx=5, pady=5)
#         select_none_but = tk.Button(but_frame, text = 'Not Paid', command =lambda: self.get_payment_data(not_paid=True), height = 1, width = 15)
#         select_none_but.grid(row=1, column=3, padx=5, pady=5)
        select_none_but = tk.Button(but_frame, text = 'Back', command = self.manage_finance, height = 1, width = 15)
        select_none_but.grid(row=2, column=1, padx=5, pady=20)
        select_none_but = tk.Button(but_frame, text = 'Start Page', command =self.start_page, height = 1, width = 15)
        select_none_but.grid(row=2, column=2, padx=5, pady=5)
    
    def get_payment_data(self,defaulter=False,save_msg=False):
        global all_comm_recs
        all_comm_recs=[]
        if defaulter:
            sql = "Select * from finance inner join member_data on finance.Id = member_data.Id where finance.actual_amt < finance.donated_amt and finance.month=? and finance.year =?"
            parameters = [datetime.date.today().strftime("%B"), datetime.date.today().strftime("%Y")]
            self.cur.execute(sql, parameters)
            data_frm_fin = self.cur.fetchall()
            sql = "Select id, donation,donation, donation=0.0, donation, name,ref_no from member_data where id not in (select id from finance where month=? and year=?)"
            parameters = [datetime.date.today().strftime("%B"), datetime.date.today().strftime("%Y")]
            self.cur.execute(sql, parameters)
            not_paid_mem = self.cur.fetchall()
            all_data=data_frm_fin+not_paid_mem
        else:
            sql = "Select * from finance inner join member_data on finance.Id = member_data.Id where finance.actual_amt >= finance.donated_amt and finance.month=? and finance.year =?"
            parameters = [datetime.date.today().strftime("%B"), datetime.date.today().strftime("%Y")]
            self.cur.execute(sql, parameters)
            all_data = self.cur.fetchall()
        self.first_frame.destroy()
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=5, column=0)
        lab_fr = tk.Frame(self.first_frame, bg=background_color)
        lab_fr.grid(row=0, column=1, padx=10,pady=1)
        but_fr = tk.Frame(self.first_frame, bg=background_color)
        but_fr.grid(row=1, column=1, padx=10,pady=1)
        self.next_but = tk.Button(but_fr, text = 'Next', height = 1, width = 12)
        self.prev_but = tk.Button(but_fr, text = 'Previous', height = 1, width = 12)
        self.next_but.grid(row=1, column=1, padx=5, pady=5)
        self.prev_but.grid(row=1, column=2, padx=5, pady=5)
        self.recs_label=tk.Label(lab_fr,bg=background_color)
        self.recs_label.grid(row=0,column=0,columnspan=5, sticky=tk.W,padx=5,pady=5)
        but_fr = tk.Frame(self.first_frame, bg=background_color)
        but_fr.grid(row=2, column=3, padx=100,pady=1,rowspan=2)
        select_all_but = tk.Button(but_fr, text = 'Select All', command = self.comm_select_all, height = 1, width = 12)
        select_all_but.grid(row=2, column=1, padx=5, pady=5)
        select_none_but = tk.Button(but_fr, text = 'Select None', command = self.comm_select_none, height = 1, width = 12)
        select_none_but.grid(row=2, column=2, padx=5, pady=5)
        prepare_but = tk.Button(but_fr, text = 'Prepare Letter', command =lambda: self.select_language(all_comm_recs,save_msg=save_msg), height = 1, width = 12)
        prepare_but.grid(row=2, column=3, padx=5, pady=5)
        back_but = tk.Button(but_fr, text = 'Back', command = lambda: self.check_current_payments(True), height = 1, width = 12)
        back_but.grid(row=3, column=2, padx=5, pady=5)
        lab_fr = tk.Frame(self.first_frame, bg=background_color)
        lab_fr.grid(row=0, column=3, padx=100,pady=1)
        self.nfd_label=tk.Label(lab_fr,bg=background_color)
        self.nfd_label.grid(row=0,column=1,columnspan=5, sticky=tk.W,padx=5)
        self.checkbox_pane = ScrollableFrame(self.master, bg=display_bg, relief=tk.RIDGE)
        self.checkbox_pane.grid(row=0, column=0, rowspan=5, columnspan=5, padx=5, pady=5)
        self.initialise_fields(payments=True)
        if not all_data:
            self.nfd_label.config(text="Database has no records to display")
            select_all_but.config(state=tk.DISABLED)
            select_none_but.config(state=tk.DISABLED)
            prepare_but.config(state=tk.DISABLED)
            self.next_but.config(state=tk.DISABLED)
            self.prev_but.config(state=tk.DISABLED)
            self.display_payments_data("", "")
        else:
            self.nfd_label.config(text="")
            self.start_num=0
            self.data_length=25
            self.max_num=25
            select_all_but.config(state=tk.NORMAL)
            select_none_but.config(state=tk.NORMAL)
            prepare_but.config(state=tk.NORMAL)
            self.max_len = len(all_data)
            if self.max_len <= self.max_num:
                self.max_num = self.max_len
                self.next_but.config(state=tk.DISABLED)
                self.prev_but.config(state=tk.DISABLED)
            for i in range(self.start_num,self.max_num,1):
                self.display_payments_data(i,all_data[i])
            self.recs_label.config( text="Displaying Records : %d to %d of %d" %(self.start_num+1,i+1,self.max_len))
            if len(all_data) > self.data_length:
                self.next_but.config(state=tk.NORMAL)
                self.next_but.config(command =lambda: self.show_next(all_data,payments=True))
                self.prev_but.config(command =lambda: self.show_previous(all_data,payments=True))
                self.prev_but.config(state=tk.DISABLED)
    
    def display_payments_data(self,ix,text):
        label0 = tk.Label(self.checkbox_pane.interior, text=datetime.date.today().strftime("%B"), width=30,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label1 = tk.Label(self.checkbox_pane.interior, text=datetime.date.today().strftime("%Y"), width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label2 = tk.Label(self.checkbox_pane.interior, text="Name", width=30,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label3 = tk.Label(self.checkbox_pane.interior, text="Ref. No", width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label4 = tk.Label(self.checkbox_pane.interior, text="Expected Amt.", width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label5 = tk.Label(self.checkbox_pane.interior, text="Actual Amt.", width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label6 = tk.Label(self.checkbox_pane.interior, text="Difference", width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label0.grid(row=0, column=0, sticky=tk.E, pady=15)
        label1.grid(row=0, column=1, sticky=tk.E)
        label2.grid(row=1, column=0, sticky=tk.E)
        label3.grid(row=1, column=1, sticky=tk.E)
        label4.grid(row=1, column=2, sticky=tk.E)
        label5.grid(row=1, column=3, sticky=tk.E)
        label6.grid(row=1, column=4, sticky=tk.E)
        if text:
            id=text[0]
            donated_amt = float(text[3])
            expected_amt = float(text[4])
            name = text[5]
            ref_no = text[6]
            self.comm_cb_v.append(tk.IntVar())
            self.comm_cb.append(tk.Checkbutton(self.checkbox_pane.interior, variable=self.comm_cb_v[ix], 
                                          command=self.comm_cb_checked, bg=display_bg))
            self.comm_cb[ix].grid(row=ix+2, column=0, sticky=tk.W, padx=5, pady=5)
            self.label_0_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %id, bg=display_bg, width=5))
            self.label_1_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %str(name), bg=display_bg, wraplength=150))
            self.label_2_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %str(expected_amt), bg=display_bg))
            self.label_3_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %str(donated_amt), bg=display_bg))
            self.label_4_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %str(donated_amt-expected_amt), bg=display_bg))
            self.label_5_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %str(ref_no), bg=display_bg))
            self.label_1_list[ix].grid(row=ix+2, column=0)
            self.label_2_list[ix].grid(row=ix+2, column=2)
            self.label_3_list[ix].grid(row=ix+2, column=3)
            self.label_4_list[ix].grid(row=ix+2, column=4)
            self.label_5_list[ix].grid(row=ix+2, column=1)
            self.dict_for[self.comm_cb[ix]] = [self.label_0_list[ix], self.label_1_list[ix], self.label_2_list[ix],
                                           self.label_3_list[ix], self.label_4_list[ix], self.label_5_list[ix] ]
        
    def comm_cb_checked(self):
        global all_comm_recs
        all_comm_recs=[]
        for ix, item in enumerate(self.comm_cb):
            if self.comm_cb_v[ix].get():
                val = self.dict_for[item][0].cget("text") + ";" +  self.dict_for[item][1].cget("text") + ";" + self.dict_for[item][2].cget("text") + ";" \
                       + self.dict_for[item][3].cget("text")
                all_comm_recs.append(val)
            
    def comm_select_all(self):
        global all_comm_recs
        all_comm_recs=[]
        for item in self.comm_cb:
            item.select()
            val = self.dict_for[item][0].cget("text") + ";" +  self.dict_for[item][1].cget("text") + ";" + self.dict_for[item][2].cget("text") + ";" \
                       + self.dict_for[item][3].cget("text")
            all_comm_recs.append(val)
            
    def comm_select_none(self):
        global all_comm_recs
        all_comm_recs=[]
        for item in self.comm_cb:
            item.deselect()
    
    def get_directory_path(self):
        curr_month = datetime.date.today().strftime("%B")
        curr_year = datetime.date.today().strftime("%Y")
        if os.path.exists(os.path.abspath(os.path.join(os.getcwd(),'saved documents'))):
            if os.path.exists(os.path.abspath(os.path.join(os.getcwd(),'saved documents','%s %s' %(curr_month, curr_year)))):
                dir_path = os.path.abspath(os.path.join(os.getcwd(), 'saved documents','%s %s' %(curr_month, curr_year)))
            else:
                os.makedirs(os.path.abspath(os.path.join(os.getcwd(),'saved documents', '%s %s' %(curr_month, curr_year))))
                dir_path = os.path.abspath(os.path.join(os.getcwd(), 'saved documents','%s %s' %(curr_month, curr_year)))
        else:
            os.makedirs(os.path.abspath(os.path.join(os.getcwd(),'saved documents')))
            os.makedirs(os.path.abspath(os.path.join(os.getcwd(),'saved documents', '%s %s' %(curr_month, curr_year))))
            dir_path = os.path.abspath(os.path.join(os.getcwd(), 'saved documents','%s %s' %(curr_month, curr_year)))
        return dir_path
    
    def prepare_letter(self, all_comm_recs, from_mem_page=False, save_msg=False):
        if self.selected_lang.get() == "":
            tkMessageBox.showwarning("Select Language", "Please select a language to continue!")
        else:
            self.first_frame.destroy()
            self.first_frame = tk.Frame(self.master, bg=background_color)
            self.first_frame.grid(row=5, column=0)
            self.text_wid = tk.Text(self.first_frame, wrap=tk.WORD)
            self.text_wid.grid(row=1, column=0, padx=50, pady=25,columnspan=3)
            scroll_b = tk.Scrollbar(self.first_frame, command=self.text_wid.yview)
            self.text_wid['yscrollcommand'] = scroll_b.set
            
            prog_frame = tk.Frame(self.first_frame,bg=background_color)
            prog_frame.grid(row=1, column=4, sticky=tk.W)
            self.progress = Progressbar(prog_frame, orient="vertical", length=250,
                                        mode="determinate")
            save_as_pdf_but = tk.Button(self.first_frame, text = 'Save as PDF', command =lambda: self.prepare_for_save(all_comm_recs,"pdf", save_msg), height = 1, width = 12)
            save_as_pdf_but.grid(row=8, column=0, padx=5, pady=5)
            save_as_word_but = tk.Button(self.first_frame, text = 'Save as .docx', command =lambda: self.prepare_for_save(all_comm_recs, "docx", save_msg), height = 1, width = 12)
            save_as_word_but.grid(row=8, column=1, padx=5, pady=5)
            back_but = tk.Button(self.first_frame, text = 'Back', height = 1, width = 12)
            back_but.grid(row=8, column=2, padx=5, pady=5)
            if from_mem_page:
                back_but.config(command=self.members_page)
            else:
                back_but.config(command=self.check_current_payments)
    
    def set_value(self,value):
        pass
    
    def check_template(self, format):
        if format == "pdf":
            if os.path.isfile(os.path.abspath(os.path.join(os.getcwd(),'template','masjid_template.pdf'))):
                pdf_temp=True
            else:
                pdf_temp=False
            return pdf_temp
        elif format == "docx":
            if os.path.isfile(os.path.abspath(os.path.join(os.getcwd(),'template','masjid_template.docx'))):
                word_temp=True
            else:
                word_temp=False
            return word_temp
        elif format == "report":
            if os.path.isfile(os.path.abspath(os.path.join(os.getcwd(),'template','report_template.pdf'))):
                report_temp=True
            else:
                report_temp=False
            return report_temp
    
    def prepare_for_save(self, all_rec, file_format, save_msg=False):
        self.recs_processed = 0
        self.max_recs = 0
        self.progress["value"]=0
        
        self.file_created=False
        text = self.text_wid.get("1.0", tk.END).strip("\n")
        if text == "":
            tkMessageBox.showerror("Missing Text", "Enter message to prepare letter!")
        else:
            temp_avail = self.check_template(file_format)
            if temp_avail:
                var = tkMessageBox.askyesno("Save file", "%s file(s) will be created in %s format!" %(str(len(all_rec)),file_format))
                if var:
                    self.progress.grid(row=0,sticky=tk.W)
                    dir_path = self.get_directory_path()
                    self.progress["maximum"] = len(all_rec)
                    i=0
                    for each_rec in all_rec:
                        self.progress["value"] = i+1
                        self.progress.update()
                        time.sleep(0.01)
                        values = each_rec.split(";")
                        id=values[0]
                        if len(values) == 4:
                            name=values[1]
                        else:
                            name=values[2]
                        sql="Select street, house_no, post_code, city, gender from member_data where id=?"
                        self.cur.execute(sql,[id])
                        data = self.cur.fetchone()
                        street = data[0]
                        h_no = data[1]
                        post_code = data[2]
                        city = data[3]
                        gender = data[4]
                        street_no = street + " " + str(h_no) + ","
                        post_city = str(post_code) + "  " + city
                        sql = "Select * from member_data where id=?"
                        self.cur.execute(sql,[id])
                        mem_data = self.cur.fetchone()
                        ref_no = str(mem_data[1]).zfill(4)
                        if save_msg:
                            sql = "Insert into comm_history values (?,?,?,?,?,?)"
                            parameters = [id,ref_no,name,str("{:%d.%m.%Y}".format(datetime.datetime.now())), text.strip("\n"), "Letter"]
                            try:
                                self.cur.execute(sql,parameters)
                                self.conn.commit()
                            except Exception, msg:
                                tkMessageBox.showerror("Error", "%s occurred while updating communication history! " %str(msg))
                        if self.selected_lang.get() == "English":
                            greeting = "Dear %s," %name
                        elif self.selected_lang.get() == "German" and gender == "Male":
                            greeting = "Sehr geehrte Herr %s," %name
                        elif self.selected_lang.get() == "German" and gender == "Female":
                            greeting = "Sehr geehrte Frau %s," %name
                        if gender == "Female":
                            sal_name = "Frau  "
                        elif gender == "Male":
                            sal_name = "Herrn  " 
                        if file_format == "pdf":
                            self.save_as_pdf(sal_name, name, street_no, post_city, greeting, text, dir_path,ref_no)
                        elif file_format == "docx":
                            self.save_as_word(sal_name, name, street_no, post_city, greeting, text, dir_path,ref_no)
                        i+=1
                    if self.file_created:
                        tkMessageBox.showinfo("Files Saved", "%s file(s) saved to directory %s" %(str(len(all_rec)),dir_path))
            else:
                tkMessageBox.showerror("Template Not Found", "Template is missing at path %s" %os.path.abspath(os.path.join(os.getcwd(),'template','masjid_template.%s'%format)))
                tkMessageBox.showerror("Abort", "File creation aborted")
                    
    def save_as_pdf(self,sal_name, name, street, city, greeting, text, dir_path,ref_no):
        reshaped_text = arabic_reshaper.reshape(text)
        bidi_text = get_display(reshaped_text)
        file_name = os.path.abspath(os.path.join(os.getcwd(),'template','masjid_template.pdf'))
        styles = getSampleStyleSheet()
        self.pdf_sal_name = sal_name
        self.pdf_name = name
        self.pdf_city = city
        self.pdf_street = street
        self.pdf_greeting = greeting
        doc = SimpleDocTemplate(os.path.join(dir_path,"temp.pdf"))
        Story = [Spacer(1,3.35*inch)]
        style = styles["Normal"]
        for each_para in bidi_text.split("\n"):
            ptext = '<font name=Times-Roman size=10>%s</font>' % each_para
            p = Paragraph(ptext, style)
            Story.append(p)
            Story.append(Spacer(1,9))
        doc.build(Story, onFirstPage=self.myFirstPage)
        with open((os.path.join(dir_path,"temp.pdf")), "rb") as f:
            new_pdf = PdfFileReader(f)
            existing_pdf = PdfFileReader(file(file_name, "rb"))
            output = PdfFileWriter()
            page = existing_pdf.getPage(0)
            page.mergePage(new_pdf.getPage(0))
            output.addPage(page)
            outputStream = file(os.path.join(dir_path, "%s_%s.pdf" %(name,ref_no)), "wb")
            output.write(outputStream)
            outputStream.close()
        if os.path.exists(os.path.join(dir_path,"temp.pdf")):
            os.remove(os.path.join(dir_path,"temp.pdf"))
        self.file_created=True
    
    def myFirstPage(self,can, doc):
        can.saveState()
        can.setFont("Times-Roman", 10)
        can.drawString(65,675,self.pdf_sal_name)
        can.drawString(65,660,self.pdf_name)
        can.drawString(65,645,self.pdf_street)
        can.drawString(65,630,self.pdf_city)
        can.setFont("Times-Roman", 10)
        can.drawString(350,590,'Date : %s' %str("{:%d.%m.%Y}".format(datetime.datetime.now())))
        can.setFont("Times-Roman", 10)
        can.drawString(65,535,self.pdf_greeting)
        can.restoreState()
    
    def save_as_word(self, sal_name, name, street_no, post_city, greeting, txt, dir_path,ref_no):
        file_name = os.path.abspath(os.path.join(os.getcwd(),'template','masjid_template.docx'))
        inp_doc = Document(file_name)
        tab = inp_doc.tables
        row = tab[0].rows
        cell = row[3].cells
        cell[0].text = sal_name
        cell = row[4].cells
        cell[0].text = name
        cell = row[5].cells
        cell[0].text = street_no
        cell = row[6].cells
        cell[0].text = post_city
        row = tab[1].rows
        cell = row[7].cells
        cell[0].text = "Date: %s" %str("{:%d.%m.%Y}".format(datetime.datetime.now()))
        para = inp_doc.paragraphs
        para[10].text = "%s" %(greeting)
        style = inp_doc.styles['Normal']
        font = style.font
        font.name = 'Times-Roman'
        font.size = Pt(10)
        para[10].style = inp_doc.styles["Normal"]
        reshaped_txt = arabic_reshaper.reshape(txt)
        bidi_text = get_display(reshaped_txt)
        for each_para in bidi_text.split("\n"):
            p = inp_doc.add_paragraph(each_para)
            p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            p.style = inp_doc.styles['Normal']
            font = style.font
            font.name = 'Times-Roman'
            font.size = Pt(10)
        inp_doc.save(os.path.join(dir_path, "%s_%s.docx" %(name,ref_no)))
        self.file_created=True
    
    def select_language(self, all_comm_recs,from_mem_page=False, save_msg=False):
        if len(all_comm_recs) < 1:
            tkMessageBox.showwarning("Prepare Letter", "Select at least 1 record!")
        else:
            self.checkbox_pane.destroy()
            self.first_frame.destroy()
            self.first_frame = tk.Frame(self.master, bg=background_color)
            self.first_frame.grid(row=5, column=0)
            frame = tk.Frame(self.first_frame, bg=background_color)
            frame.grid(row=5, column=0, pady=100, padx=100)
            l1=tk.Label(frame, text="Select Language --> ", bg=background_color)
            l1.grid(row=5,column=0, padx=50)
            lst1 = ['English','German']
            self.selected_lang = tk.StringVar()
            drop = tk.OptionMenu(frame,self.selected_lang,*lst1, command=self.set_value)
            drop.grid(row=5,column=1, padx=10, columnspan=5)
            self.selected_lang.set("")
            but = tk.Button(frame, text="Continue", command=lambda:self.prepare_letter(all_comm_recs, from_mem_page, save_msg), height=1, width=15)
            but.grid(row=6, column=0, pady=25)
            back_but = tk.Button(frame, text="Back", height=1, width=15)
            back_but.grid(row=6, column=1)
            if from_mem_page:
                back_but.config(command=self.members_page)
            else:
                back_but.config(command=self.check_current_payments)
    
    def update_finance_data(self, kill_checkpane=""):
        if kill_checkpane:
            self.checkbox_pane.destroy()
        self.first_frame.destroy()
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=5, column=0)
        lab_fr = tk.Frame(self.first_frame, bg=background_color)
        lab_fr.grid(row=0, column=1, padx=10,pady=1)
        but_fr = tk.Frame(self.first_frame, bg=background_color)
        but_fr.grid(row=1, column=1, padx=10,pady=1)
        self.next_but = tk.Button(but_fr, text = 'Next', height = 1, width = 12)
        self.prev_but = tk.Button(but_fr, text = 'Previous', height = 1, width = 12)
        self.next_but.grid(row=1, column=0, padx=5, pady=5)
        self.prev_but.grid(row=1, column=1, padx=5, pady=5)
        self.next_but.config(state=tk.DISABLED)
        self.prev_but.config(state=tk.DISABLED)
        self.recs_label=tk.Label(lab_fr,bg=background_color)
        self.recs_label.grid(row=0,column=0,columnspan=3, sticky=tk.W,pady=5,padx=5)
        but_fr = tk.Frame(self.first_frame, bg=background_color)
        but_fr.grid(row=2, column=3, padx=100,pady=1,rowspan=2)
        select_all_but = tk.Button(but_fr, text = 'Select All', command = self.finance_select_all, height = 1, width = 12)
        select_all_but.grid(row=2, column=0, padx=5, pady=5)
        select_none_but = tk.Button(but_fr, text = 'Select None', command = self.finance_select_none, height = 1, width = 12)
        select_none_but.grid(row=2, column=1, padx=5, pady=5)
        update_but = tk.Button(but_fr, text = 'Update', command = self.add_finance_details, height = 1, width = 12)
        update_but.grid(row=2, column=2, padx=5, pady=5)
        back_but = tk.Button(but_fr, text = 'Back', command = lambda: self.manage_finance(True), height = 1, width = 12)
        back_but.grid(row=3, column=0, padx=5, pady=5)
        lab_fr = tk.Frame(self.first_frame, bg=background_color)
        lab_fr.grid(row=0, column=3, padx=100,pady=1)
        self.nfd_label=tk.Label(lab_fr,bg=background_color)
        self.nfd_label.grid(row=0,column=0,columnspan=2, sticky=tk.N,padx=5)
        self.cur.execute("Select * from member_data")
        all_data = self.cur.fetchall()
        self.checkbox_pane = ScrollableFrame(self.master, bg=display_bg, relief=tk.RIDGE)
        self.checkbox_pane.grid(row=0, column=0, rowspan=5, columnspan=5, padx=5, pady=5)
        self.initialise_fields()
        self.start_num=0
        self.data_length=25
        self.max_num=25
        self.max_len = len(all_data)
        if not all_data:
            self.nfd_label.config(text="Database has no records to display!")
            select_all_but.config(state=tk.DISABLED)
            select_none_but.config(state=tk.DISABLED)
            update_but.config(state=tk.DISABLED)
            self.prev_but.config(state=tk.DISABLED)
            self.next_but.config(state=tk.DISABLED)
            self.show_data("","")
        else:
            self.nfd_label.config(text="")
            select_all_but.config(state=tk.NORMAL)
            select_none_but.config(state=tk.NORMAL)
            update_but.config(state=tk.NORMAL)
            if self.max_len <= self.max_num:
                self.max_num = self.max_len
                self.next_but.config(state=tk.DISABLED)
                self.prev_but.config(state=tk.DISABLED)
            for i in range(self.start_num,self.max_num,1):
                self.show_data(i,all_data[i])
            self.recs_label.config( text="Displaying Records : %d to %d of %d" %(self.start_num+1,i+1,self.max_len))
            if len(all_data) > self.data_length:
                self.next_but.config(state=tk.NORMAL)
                self.prev_but.config(state=tk.DISABLED)
                self.next_but.config(command=lambda: self.show_next(all_data))
                self.prev_but.config(command =lambda: self.show_previous(all_data))
    
    def initialise_fields(self,display=False,payments=False,comm_history=False):
        if display:
            self.cb = []
            self.cb_v = []
            self.label_0_list = []
            self.label_1_list = []
            self.label_2_list = []
            self.label_3_list = []
            self.label_4_list = []
            self.label_5_list = []
            self.label_6_list = []
            self.label_7_list = []
            self.label_8_list = []
            self.label_9_list = []
            self.label_10_list = []
            self.label_11_list = []
            self.label_12_list = []
            self.dict_for = {}
        elif payments:
            self.comm_cb = []
            self.comm_cb_v = []
            self.label_0_list = []
            self.label_1_list = []
            self.label_2_list = []
            self.label_3_list = []
            self.label_4_list = []
            self.label_5_list = []
            self.dict_for = {}
        elif comm_history:
            self.comm_hist_cb = []
            self.comm_hist_cb_v = []
            self.label_1_list = []
            self.label_2_list = []
            self.dict_for = {}
            self.comm_hist_cb_v=tk.StringVar()
        else:
            self.finance_cb = []
            self.finance_cb_v = []
            self.label_0_list = []
            self.label_1_list = []
            self.label_2_list = []
            self.label_3_list = []
            self.entry_0_list = []
            self.entry_1_list = []
            self.entry_2_list = []
            self.month = []
            self.year = []
            self.act_don = []
            self.dict_for = {}
        
    def show_previous(self,all_data,display=False,payments=False,comm_history=False):
        self.checkbox_pane.destroy()
        self.checkbox_pane = ScrollableFrame(self.master, bg=display_bg, relief=tk.RIDGE)
        self.checkbox_pane.grid(row=0, column=0, rowspan=5, columnspan=5, padx=5, pady=5)
        if display:
            func = "self.display_recs(x,all_data[i])"
            self.initialise_fields(display=True)
        elif payments:
            func = "self.display_payments_data(x,all_data[i])"
            self.initialise_fields(payments=True)    
        elif comm_history:
            func = "self.display_comm_data(x,all_data[i])"
            self.initialise_fields(comm_history=True)
        else:
            func = "self.show_data(x,all_data[i])"
            self.initialise_fields()
        self.start_num=self.start_num-self.data_length
        if self.start_num < 0:
            self.start_num = 0
        elif self.start_num <= 0:
            self.next_but.config(state=tk.NORMAL)
            self.prev_but.config(state=tk.DISABLED)
        else:
            self.next_but.config(state=tk.NORMAL)
            self.prev_but.config(state=tk.NORMAL)
        self.max_num=self.start_num+self.data_length
        for i,x in zip(range(self.start_num,self.max_num,1),range(self.data_length)):
            exec func
        self.recs_label.config( text="Displaying Records : %d to %d of %d" %(self.start_num+1,i+1,self.max_len))
        
    def show_next(self,all_data,display=False,payments=False,comm_history=False):
        self.checkbox_pane.destroy()
        self.checkbox_pane = ScrollableFrame(self.master, bg=display_bg, relief=tk.RIDGE)
        self.checkbox_pane.grid(row=0, column=0, rowspan=5, columnspan=5, padx=5, pady=5)
        if display:
            func = "self.display_recs(x,all_data[i])"
            self.initialise_fields(display=True)
        elif payments:
            func = "self.display_payments_data(x,all_data[i])"
            self.initialise_fields(payments=True)
        elif comm_history:
            func = "self.display_comm_data(x,all_data[i])"
            self.initialise_fields(comm_history=True)
        else:
            func = "self.show_data(x,all_data[i])"
            self.initialise_fields()
        self.start_num=self.max_num
        self.max_num=self.start_num+self.data_length
        if self.max_num >= self.max_len:
            self.max_num = self.max_len
            self.next_but.config(state=tk.DISABLED)
            self.prev_but.config(state=tk.NORMAL)
        else:
            self.prev_but.config(state=tk.NORMAL)
            self.next_but.config(state=tk.NORMAL)
        for i,x in zip(range(self.start_num,self.max_num,1),range(self.data_length)):
            exec func
        self.recs_label.config( text="Displaying Records : %d to %d of %d" %(self.start_num+1,i+1,self.max_len))
        
    def show_data(self,ix,text):  
        label0 = tk.Label(self.checkbox_pane.interior,width=5, borderwidth=1, relief=tk.SUNKEN)
        label1 = tk.Label(self.checkbox_pane.interior, text="Ref. No", width=8,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label2 = tk.Label(self.checkbox_pane.interior, text="Name", width=25,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label3 = tk.Label(self.checkbox_pane.interior, text="Expected Amt.", width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label4 = tk.Label(self.checkbox_pane.interior, text="Actual Amt.", width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label5 = tk.Label(self.checkbox_pane.interior, text="Month", width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label6 = tk.Label(self.checkbox_pane.interior, text="Year", width=6,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label0.grid(row=0, column=0, sticky=tk.E)
        label1.grid(row=0, column=1, sticky=tk.E)
        label2.grid(row=0, column=2, sticky=tk.E)
        label3.grid(row=0, column=3, sticky=tk.E)
        label4.grid(row=0, column=4, sticky=tk.E)
        label5.grid(row=0, column=5, sticky=tk.E)
        label6.grid(row=0, column=6, sticky=tk.E)
        if text:
            id=text[0]
            ref_no = text[1]
            name = text[2]
            donation = float(text[9])
            self.finance_cb_v.append(tk.IntVar())
            self.month.append(tk.StringVar(self.checkbox_pane.interior, value=datetime.date.today().strftime("%B")))
            self.year.append(tk.StringVar(self.checkbox_pane.interior, value=datetime.date.today().strftime("%Y")))
            self.act_don.append(tk.DoubleVar(self.checkbox_pane.interior, value=donation))
            try:
                self.finance_cb.append(tk.Checkbutton(self.checkbox_pane.interior, variable=self.finance_cb_v[ix], 
                                          command=self.finance_cb_checked, bg=display_bg))
                self.finance_cb[ix].grid(row=ix+1, column=0, sticky=tk.W, padx=5, pady=5)
                self.label_0_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %id, bg=display_bg))
                self.label_1_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %str(ref_no).zfill(4), bg=display_bg, width=4))
                self.label_2_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %name, bg=display_bg, wraplength=150))
                self.entry_1_list.append(tk.Entry(self.checkbox_pane.interior, width=10, textvariable=self.month[ix] ))
                self.entry_2_list.append(tk.Entry(self.checkbox_pane.interior, width=5, textvariable=self.year[ix]))
                self.entry_0_list.append(tk.Entry(self.checkbox_pane.interior, width=5, textvariable=self.act_don[ix]))
                self.label_3_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %str(donation), bg=display_bg))
                self.label_1_list[ix].grid(row=ix+1, column=1)
                self.label_2_list[ix].grid(row=ix+1, column=2)
                self.label_3_list[ix].grid(row=ix+1, column=3)
                self.entry_0_list[ix].grid(row=ix+1, column=4)
                self.entry_1_list[ix].grid(row=ix+1, column=5)
                self.entry_2_list[ix].grid(row=ix+1, column=6)
                self.dict_for[self.finance_cb[ix]] = [self.label_0_list[ix], self.label_1_list[ix], self.label_2_list[ix], self.entry_0_list[ix],
                                          self.entry_1_list[ix], self.entry_2_list[ix], self.label_3_list[ix] ]
            except Exception, msg:
                tkMessageBox.showerror("Error", str(msg))
        
    def finance_cb_checked(self):
        global all_finance_update_recs
        all_finance_update_recs=[]
        for ix, item in enumerate(self.finance_cb):
            if self.finance_cb_v[ix].get():
                val = self.dict_for[item][0].cget("text") + ";" +  self.dict_for[item][2].cget("text") + ";" + self.dict_for[item][6].cget("text") + ";" \
                        + self.dict_for[item][3].get() + ";" + self.dict_for[item][4].get() + ";" + self.dict_for[item][5].get() + ";" +\
                        str(self.finance_cb.index(item)) + ";" + self.dict_for[item][1].cget("text")
                all_finance_update_recs.append(val)
            
    def finance_select_all(self):
        global all_finance_update_recs
        all_finance_update_recs=[]
        for item in self.finance_cb:
            item.select()
            val = self.dict_for[item][0].cget("text") + ";" +  self.dict_for[item][2].cget("text") + ";" + self.dict_for[item][6].cget("text") + ";" \
                + self.dict_for[item][3].get() + ";" + self.dict_for[item][4].get() + ";" + self.dict_for[item][5].get() + ";" +\
                        str(self.finance_cb.index(item)) + ";" + self.dict_for[item][1].cget("text")
            all_finance_update_recs.append(val)
            
    def finance_select_none(self):
        global all_finance_update_recs
        all_finance_update_recs=[]
        for item in self.finance_cb:
            item.deselect()
    
    def add_finance_details(self):
        global all_finance_update_recs
        record_updated = 0
        invalid_month = False
        invalid_amt = False
        invalid_year = False
        invalid_rec_count = 0
        if len(all_finance_update_recs) < 1:
            tkMessageBox.showwarning("Update", "Select at least 1 record to update!")
        else:
            error_msg=""
            valid_months = ["january", "february", "march", "april", "may", "june", "july", "august", "september", "october", "november", "december"]
            valid_years = [year for year in range(2016,2100,1)]
            for ix, each_rec in enumerate(all_finance_update_recs):
                id,name,expected_amt,actual_amt,month,year,index_l,ref_no = each_rec.split(";")
                index_l = int(index_l)
                actual_month = self.entry_1_list[index_l].get()
                actual_amount = self.entry_0_list[index_l].get()
                actual_year = self.entry_2_list[index_l].get()
                
                if actual_month.encode("utf-8").lower() not in valid_months or actual_month == "":
                    self.entry_1_list[index_l].config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                    invalid_month = True
                    invalid_rec_count+=1
                    error_msg += "Enter valid value for month highlighted in red\n"
                else:
                    self.entry_1_list[index_l].config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                    invalid_month = False
                if actual_amount == "" :
                    self.entry_0_list[index_l].config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                    invalid_amt = True
                    invalid_rec_count+=1
                    error_msg += "Enter valid value for Actual Amt. highlighted in red\n"
                else:
                    try:
                        actual_amount = float(actual_amount)
                    except Exception, msg:
                        self.entry_0_list[index_l].config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                        invalid_amt = True
                        invalid_rec_count+=1
                        error_msg += "Enter valid value for Actual Amt. highlighted in red\n"
                    else:
                        self.entry_0_list[index_l].config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                        invalid_amt = False
                if not actual_year.isdigit():
                    self.entry_2_list[index_l].config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                    invalid_year = True
                    invalid_rec_count+=1
                    error_msg += "Enter valid value for year highlighted in red\n"
                elif int(actual_year) not in valid_years or actual_year == "":
                    self.entry_2_list[index_l].config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                    invalid_year = True
                    invalid_rec_count+=1
                    error_msg += "Enter valid value for year highlighted in red\n"
                else:
                    self.entry_2_list[index_l].config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                    invalid_year = False
                    
                if not invalid_month and not invalid_amt and not invalid_year:
                    sql = "Select * from finance where id=? and month=? and year=?"
                    self.cur.execute(sql, [id,string.capwords(actual_month),actual_year])
                    result = self.cur.fetchall()
                    if result:
                        error_msg += "Record exists for %s with Ref. No: %s for %s %s\n" %(name,ref_no,string.capwords(actual_month),actual_year) 
                    else:
                        sql = """
                                      INSERT into finance VALUES (?,?,?,?,?,?,?)
                                 """
                        parameters = [id,string.capwords(actual_month),actual_year,actual_amount,expected_amt,name,ref_no]
                        try:
                            self.cur.execute(sql, parameters)
                            self.conn.commit()
                        except Exception, msg:
                            tkMessageBox.showerror("Error", "%s occurred while deleting record! " %str(msg))
                        else:
                            record_updated += 1
            if error_msg:
                tkMessageBox.showwarning("Update",error_msg)
            tkMessageBox.showinfo("Update", "%d record(s) were updated!" %record_updated)
        self.finance_select_none()
    
    def view_all_history(self, kill_check=""):
        if kill_check:
            self.checkbox_pane.destroy()
        self.first_frame.destroy()
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=5, column=0)
        but_frame = tk.Frame(self.first_frame,bg=background_color)
        but_frame.grid(row=1,column=1, padx=100, pady=100)
        select_all_but = tk.Button(but_frame, text = 'Communication', command = self.comm_with_member, height = 1, width = 15)
        select_all_but.grid(row=1, column=1, padx=25, pady=5)
        select_none_but = tk.Button(but_frame, text = 'Deleted Members', command =self.view_deleted_members, height = 1, width = 15)
        select_none_but.grid(row=1, column=3, padx=25, pady=15)
        b2 = tk.Button(but_frame, text="Payment", command=self.finance_history, height=1, width=15)
        b2.grid(row=1, column=2, padx=5, pady=25)
        select_none_but = tk.Button(but_frame, text = 'Back', command =self.start_page, height = 1, width = 15)
        select_none_but.grid(row=2, column=1, padx=5, pady=15)
        
    def view_deleted_members(self):
        global all_member_data
        all_member_data=[]
        self.first_frame.destroy()
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=5, column=0)
        self.cur.execute("Select * from deleted_member_data")
        all_deleted_data = self.cur.fetchall()
        search_all_footer = tk.Label(self.first_frame,bg=background_color)
        search_all_footer.grid(row=8,column=0, columnspan=2,pady=5)
        self.next_but = tk.Button(self.first_frame, text = 'Next', height = 1, width = 12)
        self.prev_but = tk.Button(self.first_frame, text = 'Previous', height = 1, width = 12)
        self.next_but.grid(row=9, column=0, padx=5, pady=1)
        self.prev_but.grid(row=9, column=1, padx=5, pady=1)
        self.next_but.config(state=tk.DISABLED)
        self.prev_but.config(state=tk.DISABLED)
        self.recs_label=tk.Label(self.first_frame,bg=background_color,justify=tk.LEFT)
        self.recs_label.grid(row=8,column=0,columnspan=2, sticky=tk.W)
        back_but = tk.Button(self.first_frame, text="Back", command =lambda: self.view_all_history(True), width=15, height=1)
        back_but.grid(row=10,column=2,pady=5,padx=100)
        rest_but = tk.Button(self.first_frame, text="Restore Selected", command =self.restore_deleted_rec, width=15, height=1)
        rest_but.grid(row=9,column=2,pady=5,padx=5)
        self.checkbox_pane = ScrollableFrame(self.master, bg=display_bg, relief=tk.RIDGE)
        self.checkbox_pane.grid(row=0, column=0, rowspan=5, columnspan=5, padx=5, pady=5)
        if all_deleted_data:
            rest_but.config(state=tk.NORMAL)
            search_all_footer.config(text="")
            self.initialise_fields(display=True)
            self.start_num=0
            self.data_length=25
            self.max_num=25
            if len(all_deleted_data) < self.max_num:
                self.max_num=len(all_deleted_data)
            self.max_len = len(all_deleted_data)
            for i in range(self.start_num,self.max_num,1):
                self.display_recs(i,all_deleted_data[i])
            self.recs_label.config( text="Displaying Records : %d to %d of %d" %(self.start_num+1,i+1,self.max_len))
            if len(all_deleted_data) > self.data_length:
                self.next_but.config(state=tk.NORMAL)
                self.next_but.config(command=lambda: self.show_next(all_deleted_data, display=True))
                self.prev_but.config(command =lambda: self.show_previous(all_deleted_data, display=True))
        else:
            rest_but.config(state=tk.DISABLED)
            self.display_recs("","")
            search_all_footer.config(text="No data found for deleted members!")
    
    def generate_financial_report(self):
        self.first_frame.destroy()
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=5, column=0)
        dis_frame = tk.Frame(self.first_frame, bg=background_color)
        dis_frame.grid(row=1,column=1, padx=100, pady=150,columnspan=5)
        l1= tk.Label(dis_frame, text="Select month :", bg=background_color)
        l1.grid(row=5,column=3,sticky=tk.W)
        vals = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
        self.selected_month = tk.StringVar()
        self.select_month = tk.OptionMenu(dis_frame,self.selected_month,*vals, command=self.set_value)
        self.select_month.grid(row=5,column=4,sticky=tk.W, padx=5)
        self.selected_month.set(datetime.date.today().strftime("%B"))
        self.select_month.config(width=12)
        l1= tk.Label(dis_frame, text="Select year :", bg=background_color)
        l1.grid(row=5,column=5, padx=25,sticky=tk.W)
        vals = [year for year in range(2016,2051,1)]
        self.selected_year = tk.StringVar()
        self.select_year = tk.OptionMenu(dis_frame,self.selected_year,*vals, command=self.set_value)
        self.select_year.grid(row=5,column=6,sticky=tk.W,padx=5)
        self.selected_year.set(datetime.date.today().strftime("%Y"))
        gen_rpt = tk.Button(self.first_frame, text="Generate Report", command=self.generate_report, width=15, height=1)
        gen_rpt.grid(row=6,column=2,pady=5,padx=5)
        bck_but = tk.Button(self.first_frame, text="Back", command=self.start_page, width=15, height=1)
        bck_but.grid(row=6,column=3,pady=5,sticky=tk.NW)
    
    def pdf_first_page(self, canvas, doc):
        canvas.setFont("Times-Roman", 10)
        if self.create_for_member:
            canvas.drawString(65,720,"Finance report for : %s " %(self.member_name))
            canvas.drawString(65,705,"Member Reference No. : %s " %(self.member_ref_no))
            canvas.drawString(400,720,'Report generated on : %s' %str("{:%d.%m.%Y}".format(datetime.datetime.now())))
            canvas.setLineWidth(.2)
            canvas.line(50,695,540,695)
        else:
            canvas.drawString(65,720,"Finance report for : %s %s" %(self.selected_month.get(), self.selected_year.get()))
            canvas.drawString(65,705,'Report generated on : %s' %str("{:%d.%m.%Y}".format(datetime.datetime.now())))
            canvas.setLineWidth(.2)
            canvas.line(50,695,540,695)
        page_num = canvas.getPageNumber()
        text = "Page #%s" % page_num
        canvas.drawRightString(200*mm, 20*mm, text)
        
    def later_page(self, canvas, doc):
        canvas.setFont("Times-Roman", 10)
        page_num = canvas.getPageNumber()
        text = "Page #%s" % page_num
        canvas.drawRightString(200*mm, 20*mm, text)
    
    def create_report(self,inp_data):
        dir_path = self.get_directory_path()
        self.width, self.height = letter
        self.styles = getSampleStyleSheet()
        self.doc = SimpleDocTemplate(os.path.abspath(os.path.join(dir_path,'temp.pdf')))
        self.story = [Spacer(1, 1*inch)]
        self.createLineItems(inp_data)
        self.doc.build(self.story, onFirstPage=self.pdf_first_page, onLaterPages=self.later_page)
        
        with open(os.path.abspath(os.path.join(dir_path,'temp.pdf')), "rb") as f:
            new_pdf = PdfFileReader(f)
            existing_pdf = PdfFileReader(file(os.path.abspath(os.path.join(os.getcwd(),'template','report_template.pdf')), "rb"))
            output = PdfFileWriter()
            page = existing_pdf.getPage(0)
            page.mergePage(new_pdf.getPage(0))
            output.addPage(page)
            for page in range(new_pdf.getNumPages()-1):
                output.addPage(new_pdf.getPage(page+1))
            if self.create_for_member:
                outputStream = file(os.path.abspath(os.path.join(dir_path,"Finance_Report_%s_%s.pdf" %(self.member_name,self.member_ref_no))), "wb")
                tkMessageBox.showinfo("Report Created", "Report created at path: %s" %(os.path.abspath(os.path.join(dir_path,"Finance_Report_%s_%s.pdf"
                                                                            %(self.member_name,self.member_ref_no)))))
            else:
                outputStream = file(os.path.abspath(os.path.join(dir_path,"Finance_Report_%s_%s.pdf" %(self.selected_month.get(), self.selected_year.get()))), "wb")
                tkMessageBox.showinfo("Report Created", "Report created at path: %s" %(os.path.abspath(os.path.join(dir_path,"Finance_Report_%s_%s.pdf"
                                                                            %(self.selected_month.get(), self.selected_year.get())))))
            output.write(outputStream)
            outputStream.close()
        if os.path.isfile((os.path.abspath(os.path.join(dir_path,'temp.pdf')))):
            os.remove(os.path.abspath(os.path.join(dir_path,'temp.pdf')))
        
    def createLineItems(self,inp_data):
        if self.create_for_member:
            text_data = ["Sr. No", "Month", "Year", "Expected Amt.", "Paid Amt.", "Difference"]
        else:
            text_data = ["Sr. No", "Ref. No", "Name", "Expected Amt.", "Paid Amt.", "Difference"]
        d = []
        font_size = 11
        centered = ParagraphStyle(name="centered", alignment=TA_CENTER)
        for text in text_data:
            ptext = "<font name=Times-Roman size=%s><b>%s</b></font>" % (font_size, text)
            p = Paragraph(ptext, centered)
            d.append(p)
        data = [d]
        line_num = 1
        formatted_line_data = []
        total_expected = 0
        total_paid = 0
        total_difference = 0
        for each_rec in inp_data:
            if self.create_for_member:
                line_data = [str(line_num),each_rec[1],each_rec[2],each_rec[4],each_rec[3],(float(each_rec[3]) - float(each_rec[4]))]
                total_difference = total_difference + (float(each_rec[3]) - float(each_rec[4]))
            else:
                line_data = [str(line_num),each_rec[6],each_rec[5],each_rec[4],each_rec[3],(float(each_rec[3]) - float(each_rec[4]))]
                total_difference = total_difference + (float(each_rec[3]) - float(each_rec[4]))
            total_expected = total_expected + float(each_rec[4])
            total_paid = total_paid + float(each_rec[3])
            for item in line_data:
                ptext = "<font name=Times-Roman size=%s>%s</font>" % (font_size-1, item)
                p = Paragraph(ptext, centered)
                formatted_line_data.append(p)
            data.append(formatted_line_data)
            formatted_line_data = []
            line_num += 1
        for each_item in ["","","Total",total_expected,total_paid,total_difference]:    
            ptext = "<font name=Times-Roman size=%s>%s</font>" % (font_size-1, each_item)
            p = Paragraph(ptext, centered)
            formatted_line_data.append(p)
        data.append(["","","----------------------------------------------------","------------------","------------------", "-----------------"])    
        data.append(formatted_line_data)
        table = Table(data, colWidths=[50, 50, 190, 70, 70, 70])
        self.story.append(table)
    
    def generate_report(self,mem_data="",mem_name="",ref_no=""):
        self.create_for_member=False
        self.member_name=""
        self.member_ref_no=""
        if mem_data:
            all_data = mem_data
            self.create_for_member=True
            self.member_name = mem_name
            self.member_ref_no=ref_no
        else:
            sql ="Select * from finance where month=? and year=?"
            parameters = [self.selected_month.get(),self.selected_year.get()]
            self.cur.execute(sql,parameters)
            all_data = self.cur.fetchall()
        if all_data:
            temp_avail = self.check_template("report")
            if temp_avail:
                self.create_report(all_data)
                self.create_for_member=False
            else:
                tkMessageBox.showerror("Template Missing", "Report template is not available at path: %s" 
                                       %(os.path.abspath(os.path.join(os.getcwd(),'template','report_template.pdf'))))
        else:
            tkMessageBox.showinfo("No Data Found", "No financial data found for %s %s" %(self.selected_month.get(),self.selected_year.get()))
        
    def finance_history(self, kill_checkbox=""):
        if kill_checkbox:
            self.checkbox_pane.destroy()
        self.first_frame.destroy()
        self.first_frame = tk.Frame(self.master, bg=background_color)
        self.first_frame.grid(row=5, column=0)
        field_fr = tk.Frame(self.first_frame,bg=background_color)
        field_fr.grid(row=0,column=0, padx=150,pady=75,sticky=tk.S)
        but_fr = tk.Frame(self.first_frame,bg=background_color)
        but_fr.grid(row=1,column=0, padx=5,pady=10,sticky=tk.N)
        label = tk.Label(field_fr,text="Reference No :", bg=background_color)
        label.grid(row=0,column=0)
        self.ref_no_by_user = tk.Entry(field_fr,width=12)
        self.ref_no_by_user.grid(row=0,column=1)
        view_but = tk.Button(but_fr, text = 'View History', command = self.view_history, height = 1, width = 12)
        view_but.grid(row=0, column=1, padx=5,sticky=tk.N)
        select_none_but = tk.Button(but_fr, text = 'Back', command = self.view_all_history, height = 1, width = 12)
        select_none_but.grid(row=0, column=2, padx=5,sticky=tk.N)
    
    def view_history(self):
        if self.ref_no_by_user.get() == "":
            tkMessageBox.showwarning("View History", "Reference number cannot be blank!")
        elif not self.ref_no_by_user.get().isdigit() :
            tkMessageBox.showwarning("View History", "Please enter only numbers [0-9]")
        elif len(self.ref_no_by_user.get()) > 4 :
            tkMessageBox.showwarning("View History", "Length of Reference number cannot be greater than 4")
        else:
            stmt = "Select * from finance WHERE ref_no=?"
            parameter =  [str(self.ref_no_by_user.get()).zfill(4)]
            self.cur.execute(stmt,parameter)
            all_data = self.cur.fetchall()
            if all_data:
                ref_no = all_data[0][6]
                name = all_data[0][5]
                self.first_frame.destroy()
                self.first_frame = tk.Frame(self.master, bg=background_color)
                self.first_frame.grid(row=5, column=0)
                save_as_rpt = tk.Button(self.first_frame, text="Save as PDF", command=lambda: self.generate_report(all_data,name,ref_no), width=15, height=1)
                save_as_rpt.grid(row=1, column=0, padx=25, pady=25)
                back_but = tk.Button(self.first_frame, text="Back", command=lambda: self.finance_history(True), width=15, height=1)
                back_but.grid(row=1, column=1)
                self.checkbox_pane = ScrollableFrame(self.master, bg=display_bg, relief=tk.RIDGE)
                self.checkbox_pane.grid(row=0, column=0, rowspan=5, columnspan=5, padx=5, pady=5)
                label1 = tk.Label(self.checkbox_pane.interior, text="Name", width=15,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                label2 = tk.Label(self.checkbox_pane.interior, text="Month", width=15,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                label3 = tk.Label(self.checkbox_pane.interior, text="Year", width=20,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                label4 = tk.Label(self.checkbox_pane.interior, text="Expected Amt.", width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                label5 = tk.Label(self.checkbox_pane.interior, text="Paid Amt", width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                label6 = tk.Label(self.checkbox_pane.interior, text="Difference", width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                label1.grid(row=0, column=0, sticky=tk.E, pady=15)
                label2.grid(row=1, column=0, sticky=tk.E)
                label3.grid(row=1, column=1, sticky=tk.E)
                label4.grid(row=1, column=2, sticky=tk.E)
                label5.grid(row=1, column=3, sticky=tk.E)
                label6.grid(row=1, column=4, sticky=tk.E)
                for ix, each_rec in enumerate(all_data):
                    label1 = tk.Label(self.checkbox_pane.interior, text=each_rec[5], width=20,wraplength=150,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                    label2 = tk.Label(self.checkbox_pane.interior, text=each_rec[1], width=15,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                    label3 = tk.Label(self.checkbox_pane.interior, text=each_rec[2], width=20,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                    label4 = tk.Label(self.checkbox_pane.interior, text=each_rec[4], width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                    label5 = tk.Label(self.checkbox_pane.interior, text=each_rec[3], width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                    label6 = tk.Label(self.checkbox_pane.interior, text=float(each_rec[3]-each_rec[4]), width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
                    label1.grid(row=0, column=1, sticky=tk.E)
                    label2.grid(row=ix+2, column=0, sticky=tk.E)
                    label3.grid(row=ix+2, column=1, sticky=tk.E)
                    label4.grid(row=ix+2, column=2, sticky=tk.E)
                    label5.grid(row=ix+2, column=3, sticky=tk.E)
                    label6.grid(row=ix+2, column=4, sticky=tk.E)
            else:
                tkMessageBox.showinfo("View History", "No History found for member with reference no. %s " %str(self.ref_no_by_user.get()).zfill(4))
        
    def history_cb_checked(self):
        global all_finance_history_recs
        all_finance_history_recs=[]
        all_finance_history_recs.append(self.history_cb_v.get())
        
    def export_data(self):
        rec_format= "%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s\n"
        data_to_write = "No,Name,Street,House No,Post Code,City,Phone,Date-of-Birth,Contribution,Email,Date Joined,Comment,Gender\n"
        csv_file = tkFileDialog.asksaveasfile(mode="w", defaultextension=".csv")
        if csv_file:
                self.cur.execute("SELECT * FROM member_data")
                all_data = self.cur.fetchall()
                for each_rec in all_data:
                    data_to_write+= rec_format %(str(each_rec[1]).zfill(4), each_rec[2].encode("utf-8"), each_rec[3].encode("utf-8"), str(each_rec[4]),
                                                 str(each_rec[5]), each_rec[6].encode("utf-8"), each_rec[7],
                                                 each_rec[8].encode("utf-8"), each_rec[9], each_rec[10].encode("utf-8"), 
                                                 each_rec[11].encode("utf-8"), each_rec[12].encode("utf-8"),each_rec[13].encode("utf-8"))
                try:
                    csv_file.write(data_to_write)
                except Exception, msg:
                    tkMessageBox.showerror("Error", "%s occurred while saving file! " %str(msg))
                else:
                    csv_file.close()
                    tkMessageBox.showinfo("Export CSV", "File saved successfully!")
        else:
            tkMessageBox.showinfo("Export CSV", "Export to CSV cancelled!")

    def exit_prog(self):
        if tkMessageBox.askokcancel("Quit", "Do you want to quit?"):
            if self.conn:
                self.conn.close()
            root.destroy()
            
    def create_input_fields(self, field_list="", from_edit="", import_data=""):
        self.ref_no = tk.Entry(self.first_frame, width = 10)
        self.f_name = tk.Entry(self.first_frame, width = 30)
        self.street = tk.Entry(self.first_frame, width = 20)
        self.house_no = tk.Entry(self.first_frame, width = 6)
        self.post_code = tk.Entry(self.first_frame, width = 10)
        self.city = tk.Entry(self.first_frame, width = 30)
        self.phone = tk.Entry(self.first_frame, width = 15)
        self.date_of_birth = tk.Entry(self.first_frame, width = 12)
        self.donation = tk.Entry(self.first_frame, width = 5)
        self.email = tk.Entry(self.first_frame, width = 30)
        self.date_joined = tk.Entry(self.first_frame, width = 12)
        self.comment = tk.Text(self.first_frame, height=2, width=25, wrap=tk.WORD, font = "Arial 8")
        self.gender_entry = tk.Entry(self.first_frame, width = 10)
        vals = ["Male", "Female"]
        self.selected_gender = tk.StringVar()
        self.gender = tk.OptionMenu(self.first_frame,self.selected_gender,*vals, command=self.set_value)
        if not import_data:
            labels=[]
            for ix, text in enumerate(field_list):
                labels.append(tk.Label(self.first_frame, text=text, bg=background_color))
            labels[0].grid(row=0, column=1, sticky=tk.E, padx=1, pady=10)
            labels[1].grid(row=0, column=3, sticky=tk.E, padx=1, pady=10)
            labels[2].grid(row=1, column=1, sticky=tk.E, padx=1, pady=10)
            labels[3].grid(row=1, column=3, sticky=tk.E, padx=1, pady=10)
            labels[4].grid(row=2, column=1, sticky=tk.E, padx=1, pady=10)
            labels[5].grid(row=2, column=3, sticky=tk.E, padx=1, pady=10)
            labels[6].grid(row=3, column=1, sticky=tk.E, padx=1, pady=10)
            labels[7].grid(row=3, column=3, sticky=tk.E, padx=1, pady=10)
            labels[8].grid(row=4, column=1, sticky=tk.E, padx=1, pady=10)
            labels[9].grid(row=4, column=3, sticky=tk.E, padx=1, pady=10)
            labels[10].grid(row=5, column=1, sticky=tk.E, padx=1, pady=10)
            labels[11].grid(row=5, column=3, sticky=tk.E, padx=1, pady=10)
            labels[12].grid(row=6, column=1, sticky=tk.E, padx=1, pady=10)
            self.ref_no.grid(row=0, column=2, sticky=tk.W)
            self.f_name.grid(row=0, column=4, sticky=tk.W, columnspan=3)
            self.street.grid(row=1, column=2, sticky=tk.W)
            self.house_no.grid(row=1, column=4, sticky=tk.W, columnspan=3)
            self.post_code.grid(row=2, column=2, sticky=tk.W)
            self.city.grid(row=2, column=4, sticky=tk.W, columnspan=3)
            self.phone.grid(row=3, column=2, sticky=tk.W)
            self.date_of_birth.grid(row=3, column=4, sticky=tk.W)
            self.donation.grid(row=4, column=2, sticky=tk.W)
            self.email.grid(row=4, column=4, sticky=tk.W, columnspan=3)
            self.date_joined.grid(row=5, column=2, sticky=tk.W)
            self.comment.grid(row=6, column=2, sticky=tk.W, columnspan=5, rowspan=3, pady=5)
            if from_edit:
                self.gender_entry = tk.Entry(self.first_frame, width = 8)
                self.gender_entry.grid(row=5,column=4,sticky=tk.W)
            else:    
                self.gender.grid(row=5,column=4,sticky=tk.W, columnspan=3)
                self.selected_gender.set("")
            self.footer = tk.Label(self.first_frame, text="", fg="red", bg=background_color)
            self.footer.grid(row=18, column=1, columnspan=10, pady=10)
    
    def validate_input(self,edit_rec="",from_import=""):
        name = self.f_name.get().replace(" ", "")
        street = self.street.get().replace(" ", "")
        city = self.city.get().replace(" ", "")
        email_validated = True
        error_msg = ""
        if not edit_rec:
            if not self.ref_no.get().isdigit():
                self.ref_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                error_msg += "Please enter only numbers[0-9] for Reference Number! \n"
                ref_no_verified = False
            elif len(self.ref_no.get().zfill(4)) > 4:
                self.ref_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                error_msg += "Please enter only 4 numbers[1234] for Reference Number! \n"
                ref_no_verified = False
            else:
                stmt = "select * from member_data where ref_no=?"
                self.cur.execute(stmt, [str(self.ref_no.get()).zfill(4)])
                ref_no = self.cur.fetchone()
                
                stmt = "select * from deleted_member_data where ref_no=?"
                self.cur.execute(stmt, [str(self.ref_no.get()).zfill(4)])
                del_ref_no = self.cur.fetchone()
                
                if ref_no or del_ref_no:
                    self.ref_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                    error_msg += "Reference number: %s is already assigned to a member. Please use another number! \n" %str(self.ref_no.get().zfill(4))
                    ref_no_verified = False
                else:
                    self.ref_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                    ref_no_verified = True
        else:
            ref_no_verified = True
            
        if not name.isalpha():
            self.f_name.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            error_msg += "Please enter only alphabets[a-z|A-Z] for Name! \n"
            f_name_verified = False
        else:
            self.f_name.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            f_name_verified = True
        
        if self.donation.get() !="":
            try:
                donation = float(self.donation.get())
            except Exception, msg:
                error_msg += "Please enter valid amount for contribution!\n"
                self.donation.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                donation_verified = False
            else:
                self.donation.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                donation_verified = True
        else:
            error_msg += "Please enter valid amount for contribution!\n"
            donation_verified = False
        
        if self.date_of_birth.get():
            date_of_birth_verified,msg = self.validate_date_fields(self.date_of_birth.get())
            if not date_of_birth_verified:
                error_msg += msg + " for Date of Birth!\n"
                self.date_of_birth.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            else:
                self.date_of_birth.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
        else:
            error_msg += "Date of birth is missing!"
            date_of_birth_verified=False
        
        if not from_import:    
            
            if not street.isalpha():
                error_msg += "Please enter only alphabets[a-z|A-Z] for Street!\n"
                self.street.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                street_verified = False
            else:
                self.street.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                street_verified = True
                
            if not self.house_no.get().isalnum():
                error_msg += "Please enter only numbers[0-9] for House No!\n"
                self.house_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                house_no_verified = False
            else:
                self.house_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                house_no_verified = True
                
            if not self.post_code.get().isdigit():
                error_msg += "Please enter only numbers[0-9] for Post Code!\n"
                self.post_code.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                post_code_verified = False
            else:
                self.post_code.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                post_code_verified = True   
                 
            if not city.isalpha():
                error_msg += "Please enter only alphabets[a-z|A-Z] for City!\n"
                self.city.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                city_verified = False
            else:
                self.city.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                city_verified = True
                
            if not self.phone.get().isdigit():
                error_msg += "Please enter only numbers[0-9] for Phone Number!\n"
                self.phone.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                phone_verified = False
            else:
                self.phone.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                phone_verified = True
                
            if self.email.get():
                if not self.email.index("end") == 0:
                    email_format =  "^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$"
                    if not re.match(email_format, self.email.get()):
                        self.email.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                        error_msg += "Enter email in the format example@email.com\n"
                        email_validated = False
                    else:
                        self.email.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                        email_validated = True
                else:
                    self.email.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                    email_validated = True
            else:
                self.email.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
                email_validated = True
            
    
            
            if self.date_joined.get():
                date_joined_verified,msg = self.validate_date_fields(self.date_joined.get())
                if not date_joined_verified:
                    error_msg += msg + " for Date of Joining!\n"
                    self.date_joined.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
                else:
                    self.date_joined.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            else:
                error_msg += "Date of joining is missing!"
                date_joined_verified=False
                
        else:
            street_verified = house_no_verified = post_code_verified = city_verified = phone_verified = email_validated = date_joined_verified = True   
             
            
        if (f_name_verified and street_verified and house_no_verified and post_code_verified and city_verified and phone_verified and  
            donation_verified and email_validated and ref_no_verified and date_joined_verified and date_of_birth_verified):
            self.all_fields_verified = True
        else:
            if from_import:
                return error_msg
            else:
                tkMessageBox.showwarning("Invalid Input", error_msg)
            
    def check_empty_fields(self, from_edit=""):
        self.no_empty_field = False
        if from_edit:
            gender_verified=True
        else:
            if self.selected_gender.get() == "":
                gender_verified = False
                self.gender.config(bg="red")
            else:
                gender_verified = True
                self.gender.config(bg="white")
        if self.ref_no.index("end") == 0:
            self.ref_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            ref_no_not_empty = False
        else:
            self.ref_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            ref_no_not_empty = True
        if self.f_name.index("end") == 0:
            self.f_name.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            f_name_not_empty = False
        else:
            self.f_name.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            f_name_not_empty = True
        if self.street.index("end") == 0:
            self.street.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            street_not_empty = False
        else:
            self.street.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            street_not_empty = True
        if self.house_no.index("end") == 0:
            self.house_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            house_no_not_empty = False
        else:
            self.house_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            house_no_not_empty = True    
        if self.post_code.index("end") == 0:
            self.post_code.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            post_code_not_empty = False
        else:
            self.post_code.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)    
            post_code_not_empty = True
        if self.city.index("end") == 0:
            self.city.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            city_not_empty = False
        else:
            self.city.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            city_not_empty=True
        if self.phone.index("end") == 0:
            self.phone.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            phone_not_empty = False
        else:
            self.phone.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            phone_not_empty = True
        if self.donation.index("end") == 0:
            self.donation.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            donation_not_empty = False
        else:
            self.donation.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            donation_not_empty = True
        if self.date_of_birth.index("end") == 0:
            self.date_of_birth.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            date_of_birth_not_empty = False
        else:
            self.date_of_birth.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            date_of_birth_not_empty = True
        if self.date_joined.index("end") == 0:
            self.date_joined.config(highlightcolor="red", highlightbackground="red", highlightthickness=1)
            date_joined_not_empty = False
        else:
            self.date_joined.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            date_joined_not_empty = True
            
        if gender_verified and ref_no_not_empty and f_name_not_empty and street_not_empty and post_code_not_empty and house_no_not_empty and city_not_empty and phone_not_empty \
            and donation_not_empty and date_of_birth_not_empty and date_joined_not_empty:
            self.no_empty_field = True
        else:
            tkMessageBox.showwarning("Invalid Input", "Please enter value for field(s) highlighted in Red!")
                
    def clear_input_fields(self,from_edit=""):
        self.ref_no.delete(0, tk.END)
        self.f_name.delete(0, tk.END)
        self.street.delete(0, tk.END)
        self.house_no.delete(0, tk.END)
        self.post_code.delete(0, tk.END)
        self.city.delete(0, tk.END)
        self.phone.delete(0, tk.END)
        self.date_of_birth.delete(0, tk.END)
        self.email.delete(0, tk.END)
        self.donation.delete(0, tk.END)
        self.comment.delete("1.0", tk.END)
        self.date_joined.delete(0, tk.END)
        if not from_edit:
            self.selected_gender.set("")
                
    def save_record(self, id):
        self.all_fields_verified = False
        self.check_empty_fields(True)
        if self.no_empty_field :
            self.validate_input(edit_rec=True)
        if self.all_fields_verified and self.no_empty_field:
            ans = tkMessageBox.askyesno("Edit", "Are you sure you want to save the changes?")
            if ans:
                if self.donation.get() == "":
                    donation = 0
                else:
                    donation = float(self.donation.get())
                    
                sql = """ UPDATE member_data
                                    SET name=?, street=?, house_no=?, post_code=?, city=?, phone=?, date_of_birth=?, email=?, donation=? , 
                                     date_joined=?, comment=?
                                    WHERE Id = ? """
                parameters = [self.f_name.get(), self.street.get(), self.house_no.get(), int(self.post_code.get()), self.city.get(), 
                                    int(self.phone.get()), self.date_of_birth.get(), self.email.get(), donation, self.date_joined.get(),
                                     self.comment.get("1.0",tk.END).strip("\n"), id]
                try:
                    self.cur.execute(sql,parameters)
                    self.conn.commit()
                except Exception, msg:
                    tkMessageBox.showerror("Error", "%s occurred while saving edited record! " %str(msg))
                else:    
                    tkMessageBox.showinfo("Edit", "Record changed successfully!")
                    self.clear_input_fields(True)
                    self.members_page()

    def fetch_record_and_display(self):            
        result_list = []
        self.cur.execute(""" SELECT * from member_data""")
        all_lines = self.cur.fetchall()
        self.checkbox_pane.destroy()
        for each_line in all_lines:
            ref_no = each_line[1]
            name = each_line[2]
            city = each_line[6]
            phone_num = each_line[7]
            if self.var.get() == "Name" and (self.s_name.get().lower() in name.lower()) :
                result_list.append(each_line)
                continue
            elif self.var.get() == "City" and (self.s_city.get().lower() in city.lower()) :
                result_list.append(each_line)
                continue
            elif self.var.get() == "Phone" and (self.s_phone.get() in str(phone_num)) :
                result_list.append(each_line)
                continue
            elif self.var.get() == "Ref. No" and (str(self.s_ref_no.get()).zfill(4) == ref_no.zfill(4)) :
                result_list.append(each_line)
                continue
        self.checkbox_pane = ScrollableFrame(self.master, bg=display_bg, relief=tk.RIDGE)
        self.checkbox_pane.grid(row=0, column=0, rowspan=5, columnspan=5, padx=5, pady=5)
        self.initialise_fields(display=True)
        self.recs_label.config( text="")
        if result_list:
            self.search_field_footer.config(text="")
            self.start_num=0
            self.data_length=25
            self.max_num=25
            self.max_len = len(result_list)
            if self.max_len <= self.max_num:
                self.max_num = self.max_len
                self.next_but.config(state=tk.DISABLED)
                self.prev_but.config(state=tk.DISABLED)
            for i in range(self.start_num,self.max_num,1):
                self.display_recs(i,result_list[i])
            self.recs_label.config( text="Displaying Records : %d to %d of %d" %(self.start_num+1,i+1,self.max_len))
            if len(result_list) > self.data_length:
                self.next_but.config(command=lambda: self.show_next(result_list, display=True))
                self.prev_but.config(command =lambda: self.show_previous(result_list, display=True))
        else:
            self.display_recs("","")
            self.search_field_footer.config(text="No record(s) found for search criteria" )
    
    def selected_rb(self):
        if self.var.get() == "City":
            self.s_city.config(state="normal")
            self.s_name.delete(0, 'end')
            self.s_phone.delete(0, 'end')
            self.s_ref_no.delete(0, 'end')
            self.disable_red_border_clear_text(self.var.get())
            self.s_name.config(state="disabled", disabledbackground=label_bg)
            self.s_phone.config(state="disabled", disabledbackground=label_bg)
            self.s_ref_no.config(state="disabled", disabledbackground=label_bg)
        elif self.var.get() == "Phone":
            self.s_phone.config(state="normal")
            self.s_name.delete(0, 'end')
            self.s_city.delete(0, 'end')
            self.s_ref_no.delete(0, 'end')
            self.disable_red_border_clear_text(self.var.get())
            self.s_name.config(state="disabled", disabledbackground=label_bg)
            self.s_city.config(state="disabled", disabledbackground=label_bg)
            self.s_ref_no.config(state="disabled", disabledbackground=label_bg)
        elif self.var.get() == "Name":
            self.s_name.config(state="normal")
            self.s_phone.delete(0, 'end')
            self.s_city.delete(0, 'end')
            self.s_ref_no.delete(0, 'end')
            self.disable_red_border_clear_text(self.var.get())
            self.s_phone.config(state="disabled", disabledbackground=label_bg)
            self.s_city.config(state="disabled", disabledbackground=label_bg)
            self.s_ref_no.config(state="disabled", disabledbackground=label_bg)
        elif self.var.get() == "Ref. No":
            self.s_ref_no.config(state="normal")
            self.s_name.delete(0, 'end')
            self.s_phone.delete(0, 'end')
            self.s_city.delete(0, 'end')
            self.disable_red_border_clear_text(self.var.get())
            self.s_phone.config(state="disabled", disabledbackground=label_bg)
            self.s_city.config(state="disabled", disabledbackground=label_bg)
            self.s_name.config(state="disabled", disabledbackground=label_bg)
            
    def display_recs(self,ix,text):
        label0 = tk.Label(self.checkbox_pane.interior,width=5, borderwidth=1, relief=tk.SUNKEN)
        label1 = tk.Label(self.checkbox_pane.interior, text="Ref. No", width=8,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label2 = tk.Label(self.checkbox_pane.interior, text="Name", width=25,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label3 = tk.Label(self.checkbox_pane.interior, text="Street,No", width=20,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label4 = tk.Label(self.checkbox_pane.interior, text="PostCode,City", width=20,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label5 = tk.Label(self.checkbox_pane.interior, text="Phone", width=12,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label6 = tk.Label(self.checkbox_pane.interior, text="Birth date", width=10,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label7 = tk.Label(self.checkbox_pane.interior, text="Email", width=35,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label8 = tk.Label(self.checkbox_pane.interior, text="Contribution", width=10,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label9 = tk.Label(self.checkbox_pane.interior, text="Date Joined", width=10,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label10 = tk.Label(self.checkbox_pane.interior, text="Comment", width=23,justify=tk.CENTER, borderwidth=1, relief=tk.SUNKEN)
        label0.grid(row=0, column=0, sticky=tk.E)
        label1.grid(row=0, column=1, sticky=tk.E)
        label2.grid(row=0, column=2, sticky=tk.E)
        label3.grid(row=0, column=3, sticky=tk.E)
        label4.grid(row=0, column=4, sticky=tk.E)
        label5.grid(row=0, column=5, sticky=tk.E)
        label6.grid(row=0, column=8, sticky=tk.E)
        label7.grid(row=0, column=7, sticky=tk.E)
        label8.grid(row=0, column=6, sticky=tk.E)
        label9.grid(row=0, column=9, sticky=tk.E)
        label10.grid(row=0, column=10, sticky=tk.E)
        if text:
            id=text[0]
            ref_no = str(text[1]).zfill(4)
            name = text[2]
            street = text[3]
            h_no = text[4]
            post_code = text[5]
            city = text[6]
            phone = text[7]
            date_of_birth = text[8]
            email = text[9]
            donation = text[10]
            date =  text[11]
            comment = text[12]
            gender=text[13]
            self.cb_v.append(tk.IntVar())
            self.cb.append(tk.Checkbutton(self.checkbox_pane.interior, variable=self.cb_v[ix], 
                                          command=self.cb_checked, bg=display_bg))
            self.cb[ix].grid(row=ix+1, column=0, sticky=tk.W, padx=5, pady=5)
            self.label_0_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %id, bg=display_bg))
            self.label_1_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %str(ref_no), bg=display_bg, width=4))
            self.label_2_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %name, bg=display_bg, wraplength=150))
            self.label_3_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %(street + "\n" + str(h_no)), bg=display_bg,justify=tk.CENTER ))
            self.label_4_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %(str(post_code) + "\n" + city), bg=display_bg, justify=tk.CENTER ))
            self.label_5_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %phone, bg=display_bg))
            self.label_6_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %email, bg=display_bg))
            self.label_7_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %donation, bg=display_bg))
            self.label_8_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %date_of_birth, bg=display_bg))
            self.label_9_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %date, bg=display_bg))
            self.label_10_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %comment, bg=display_bg, wraplength=150, justify=tk.CENTER))
            self.label_11_list.append(tk.Label(self.checkbox_pane.interior, text="%s" %gender, bg=display_bg))
            self.label_1_list[ix].grid(row=ix+1, column=1)
            self.label_2_list[ix].grid(row=ix+1, column=2)
            self.label_3_list[ix].grid(row=ix+1, column=3)
            self.label_4_list[ix].grid(row=ix+1, column=4)
            self.label_5_list[ix].grid(row=ix+1, column=5)
            self.label_6_list[ix].grid(row=ix+1, column=6)
            self.label_7_list[ix].grid(row=ix+1, column=7)
            self.label_8_list[ix].grid(row=ix+1, column=8)
            self.label_9_list[ix].grid(row=ix+1, column=9)
            self.label_10_list[ix].grid(row=ix+1, column=10)
            self.dict_for[self.cb[ix]] = [self.label_0_list[ix], self.label_1_list[ix], self.label_2_list[ix], self.label_3_list[ix], self.label_4_list[ix],
                                          self.label_5_list[ix], self.label_6_list[ix], self.label_7_list[ix], self.label_8_list[ix], 
                                          self.label_9_list[ix], self.label_10_list[ix], self.label_11_list[ix]]

    def cb_checked(self):
        global all_member_data
        all_member_data=[]
        for ix, item in enumerate(self.cb):
            if self.cb_v[ix].get():
                val = self.dict_for[item][0].cget("text") + ";" + self.dict_for[item][1].cget("text") + ";" + self.dict_for[item][2].cget("text") + \
                        ";" + self.dict_for[item][3].cget("text") + ";" + self.dict_for[item][4].cget("text") + ";" + \
                        self.dict_for[item][5].cget("text") + ";" + self.dict_for[item][6].cget("text") + ";" + self.dict_for[item][7].cget("text") + \
                        ";" + self.dict_for[item][8].cget("text") + ";" + self.dict_for[item][9].cget("text")+ ";" + self.dict_for[item][10].cget("text") + \
                        ";" + self.dict_for[item][11].cget("text")
                all_member_data.append(val)
    
    def restore_deleted_rec(self):
        global all_member_data
        if len(all_member_data) < 1:
            tkMessageBox.showwarning("Restore", "Select at least 1 record to restore!")
        else:
            var = tkMessageBox.askyesno("Restore", "Do you really want to restore %s record(s)?" %str(len(all_member_data)))
            if var:
                for each_value in all_member_data:
                    values = each_value.split(";")
                    old_id = values[0]
                    street,h_no = values[3].split("\n")
                    post_c,city = values[4].split("\n")
                    sql = "INSERT into member_data VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)"
                    parameters = [old_id,values[1],values[2],street,h_no,post_c,city,values[5],values[8],values[6],values[7],values[9],
                                  "Reactivated on: %s"%str("{:%d.%m.%Y}".format(datetime.datetime.now())),values[11]]
                    try:
                        self.cur.execute(sql,parameters)
                        self.conn.commit()
                    except Exception, msg:
                        tkMessageBox.showerror("Error", "%s occurred while activating deleted record! " %str(msg))
                    else:
                        try:
                            self.cur.execute("DELETE from deleted_member_data where Id = ?",[(old_id)])
                            self.conn.commit()
                        except Exception, msg:
                            tkMessageBox.showerror("Error", "%s occurred while deleting record to activate it! " %str(msg))
                tkMessageBox.showinfo("Restore", "%s record(s) restored successfully!" %str(len(all_member_data)))
                self.checkbox_pane.destroy()
                self.view_deleted_members()
            
    def disable_red_border_clear_text(self,field):
        if field == "Name":
            self.s_city.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.s_phone.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.s_ref_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.label_ref_no.config(text="")
            self.label_phone.config(text="")
            self.label_city.config(text="")
        elif field == "City":
            self.s_name.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.s_phone.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.s_ref_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.label_ref_no.config(text="")
            self.label_phone.config(text="")
            self.label_name.config(text="")
        elif field == "Phone":
            self.s_name.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.s_city.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.s_ref_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.label_ref_no.config(text="")
            self.label_name.config(text="")
            self.label_city.config(text="")
        elif field == "Ref. No":
            self.s_name.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.s_city.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.s_phone.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.label_name.config(text="")
            self.label_city.config(text="")
            self.label_phone.config(text="")#
        elif field == "All":
            self.s_city.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.s_phone.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.s_name.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.s_ref_no.config(highlightcolor="red", highlightbackground="red", highlightthickness=0)
            self.label_ref_no.config(text="")
            self.label_name.config(text="")
            self.label_city.config(text="")
            self.label_phone.config(text="")
            self.s_name.delete(0, 'end')
            self.s_city.delete(0, 'end')
            self.s_phone.delete(0, 'end')
            self.s_ref_no.delete(0, 'end')
            self.search_field_footer.config(text="")
    
    def validate_date_fields(self,date):
        error_msg=""
        valid_day_month = False
        min_year = 1900
        max_year = min_year + 200
        if len(date)!= 10:
            error_msg = "Please enter date in format DD/MM/YYYY"
        else:
            day,month,year = date.split("/")
            if len(day) != 2:
                error_msg = "Length of day in not 2. Please enter day as 01 for first!" 
            elif len(month) != 2:
                error_msg = "Length of month in not 2. Please enter month as 01 for January!"
            elif len(year) != 4:
                error_msg = "Length of year in not 4. Please enter year as 2001!"
            else:
                day=int(day)
                month=int(month)
                year=int(year)
                if day < 1 or day > 31:
                    error_msg = "Day %s is not in range [01-31]" %str(day)
                else:
                    if month < 1 or month > 12:
                        error_msg = "Month %s is not in range [01-12]" %str(month)
                    else:
                        if year < min_year or year > max_year:
                            error_msg = "Year is not in range [%s-%s]" %(str(min_year), str(max_year))
                        else:
                            if year%4 == 0:
                                leap_year = True
                            else:
                                leap_year = False
                            valid_day_month,error_msg = self.validate_day_month(day,str(month).zfill(2),leap_year)
        return valid_day_month,error_msg

    def validate_day_month(self,day,month,leap_year_or_not):
        days_31 = ['01', '03', '05', '07', '08', '10', '12']
        days_30 = ['04', '06', '10', '11']
        days_28 = ['02']
        
        month_dict = {'01':'January',
                  '02':'February',
                  '03':'March',
                  '04':'April',
                  '05':'May',
                  '06':"June",
                  '07':'July',
                  '08':'August',
                  '09':'September',
                  '10':'October',
                  '11':'November',
                  '12':'December'}
        if leap_year_or_not:
            max_day = 29
        else:
            max_day=28
        if month in days_28:
            if day > max_day:
                return False, "Invalid day: %s for month %s" %(day, month_dict[month])
            else:
                return True,""
        elif month in days_30:
            if day > 30:
                return False, "Invalid day: %s for month %s" %(day, month_dict[month])
            else:
                return True,""
        elif month in days_31:
            if day <= 31:
                return True,""
            else:
                return False, "Invalid day: %s for month %s" %(day, month_dict[month])
        else:
            return False, "Invalid month:%s, Invalid day:%s" %(month,day)
        
    def select_import_file(self):    
        file_path = tkFileDialog.askopenfilename()
        if file_path:
            self.recs_processed = 0
            self.max_recs = 0
            self.progress["value"]=0
            self.import_data_from_csv(file_path)
        else:
            tkMessageBox.showinfo("Import", "Data import cancelled by user")
        
    def import_data_from_csv(self,file_path):
        error_msg=""
        self.progress.grid(row=0,sticky=tk.W)
        with open(file_path,"rb") as csv_file:
            csv_reader = csv.reader(csv_file, delimiter=',')
            data = list(csv_reader)
            total_rec = len(data)
            self.progress["maximum"] = total_rec
            i=0
            gender = ["male", "female"]
            for each_row in data:
                self.progress["value"] = i+1
                self.progress.update()
                time.sleep(0.01)
                if i==0:
                    i=+1
                    continue
                if len(each_row) != 13:
                    error_msg += "At line %d ==> Record length %d is invalid. Should be %d\n" %(i, len(each_row),13)
                    i+=1
                    continue
                else:
                    ref_no = each_row[0]
                    if ref_no == None or ref_no=="":
                        error_msg+="At line %d ==> Reference number is missing\n" %str(i)
                        i+=1
                        continue
                    elif len(ref_no) > 4:
                        error_msg+="At line %d ==> Please enter only 4 numbers [1234] for Reference number\n" %i
                        i+=1
                        continue
                    else:
                        new_ref_no=ref_no.zfill(4)
                    try:
                        new_ref_no = int(new_ref_no)
                    except Exception, msg:
                        error_msg+= "At line %d ==> Please enter only numbers [0-9] for Reference number \n" %i
                        i+=1
                        continue
                    else:
                        name = each_row[1].strip()
                        street = each_row[2].strip()
                        h_no = each_row[3].strip()
                        post_code = each_row[4].strip()
                        city = each_row[5].decode("utf-8").strip()
                        phone = each_row[6].strip()
                        date_of_b = each_row[7].strip()
                        donation = each_row[8].strip()
                        email = each_row[9].strip()
                        date_of_j = each_row[10].strip()
                        comment = each_row[11].strip()
                        m_or_f = each_row[12].strip()
#                         if m_or_f=="":
#                             error_msg+="At line %d ==> Gender is missing \n" %i
#                             i+=1
#                             continue
#                         elif m_or_f.lower() not in gender:
#                             error_msg+="At line %d ==> Enter only Male or Female for gender \n" %i
#                             i+=1
#                             continue
#                         else:
                        self.create_input_fields(import_data=True)
                        self.selected_gender.set(m_or_f)
                        self.ref_no.insert(0, new_ref_no)
                        self.f_name.insert(0, name)
                        self.street.insert(0, street)
                        self.house_no.insert(0, h_no)
                        self.post_code.insert(0, post_code)
                        self.city.insert(0, city)
                        self.phone.insert(0, phone)
                        self.date_of_birth.insert(0, date_of_b)
                        self.email.insert(0, email)
                        self.donation.insert(0, donation)
                        self.comment.insert("1.0", comment)
                        self.date_joined.insert(0, date_of_j)
                        
                        self.no_empty_field=True
                        msg = self.add_record(from_import=True)
                        error_msg+="At line %d ==> " %i + msg
                        i+=1
                        self.clear_input_fields(False)
                        self.no_empty_field=False
                                                                    
            with open("import_log.txt", "w") as log_file:
                log_file.write(error_msg)
            tkMessageBox.showinfo("Import Data","Data import completed. Check file %s for details!" %(os.path.abspath(os.path.join(os.getcwd(),"import_log.txt"))))

        
def exit_prog():
    var = tkMessageBox.askokcancel("Quit", "Do you want to quit?")
    if var:
        root.destroy()
import tempfile
import tkFont
root = tk.Tk()

icon="""R0lGODlhUgEPAXAAACH5BAEAAPwALAAAAABSAQ8BhwAAAAAAMwAAZgAAmQAAzAAA/wArAAArMwArZgArmQArzAAr/wBVAABVMwBV
ZgBVmQBVzABV/wCAAACAMwCAZgCAmQCAzACA/wCqAACqMwCqZgCqmQCqzACq/wDVAADVMwDVZgDVmQDVzADV/wD/AAD/MwD/ZgD/
mQD/zAD//zMAADMAMzMAZjMAmTMAzDMA/zMrADMrMzMrZjMrmTMrzDMr/zNVADNVMzNVZjNVmTNVzDNV/zOAADOAMzOAZjOAmTOA
zDOA/zOqADOqMzOqZjOqmTOqzDOq/zPVADPVMzPVZjPVmTPVzDPV/zP/ADP/MzP/ZjP/mTP/zDP//2YAAGYAM2YAZmYAmWYAzGYA
/2YrAGYrM2YrZmYrmWYrzGYr/2ZVAGZVM2ZVZmZVmWZVzGZV/2aAAGaAM2aAZmaAmWaAzGaA/2aqAGaqM2aqZmaqmWaqzGaq/2bV
AGbVM2bVZmbVmWbVzGbV/2b/AGb/M2b/Zmb/mWb/zGb//5kAAJkAM5kAZpkAmZkAzJkA/5krAJkrM5krZpkrmZkrzJkr/5lVAJlV
M5lVZplVmZlVzJlV/5mAAJmAM5mAZpmAmZmAzJmA/5mqAJmqM5mqZpmqmZmqzJmq/5nVAJnVM5nVZpnVmZnVzJnV/5n/AJn/M5n/
Zpn/mZn/zJn//8wAAMwAM8wAZswAmcwAzMwA/8wrAMwrM8wrZswrmcwrzMwr/8xVAMxVM8xVZsxVmcxVzMxV/8yAAMyAM8yAZsyA
mcyAzMyA/8yqAMyqM8yqZsyqmcyqzMyq/8zVAMzVM8zVZszVmczVzMzV/8z/AMz/M8z/Zsz/mcz/zMz///8AAP8AM/8AZv8Amf8A
zP8A//8rAP8rM/8rZv8rmf8rzP8r//9VAP9VM/9VZv9Vmf9VzP9V//+AAP+AM/+AZv+Amf+AzP+A//+qAP+qM/+qZv+qmf+qzP+q
///VAP/VM//VZv/Vmf/VzP/V////AP//M///Zv//mf//zP///wAAAAAAAAAAAAAAAAj/AMkJHEiwoMGDCBMqXMiwocOHECNKnEix
osWLGC1Gqxet2LJlxTZhBKlpjRoyOmjQ0KGAQMuXLl2qZKlDB8oyatRU0rTJY8af5DZpKga0qNGjSCUuu7j0o0ehmzZRikNpDRk1
ZWqmVAkzpoIBLXVMXLaJDMyzZ10qeLlWbUyvcNsSsAmHaNKGBLzqWMPT7t2/5D5CFVoJzs6oPEF+bBqRI+CEZl3axDq16tWsNVXS
SNsyQVyvnNeyBT0XYtmWake3bbuZBgHXm12sViBbtOrUXh8n3DRbttu5Nk1CXay7IVlKNTnnDU3As3IFM6GvlA5TZPGBvOOiVVt7
tlvbrr2v/938vXZ3tWSMq9neknzrmihRTs551aZN+ZrNr00Q3u3S6wRFttp3aX22nFpczQRfTVapsQYlwzkGUTFqsKRaba/hxhpt
ohGooYcdKgCgQJu1N6B4Z4Wnoku+QSfTa1vdd1JOmlDCkybEpFTaQjqAJp0ONhajDDOdEGlkkUgeqWSSTC6jTEg56dCiJiMOVOJK
r5n4XYkZJmDbi18yd+JyK5Hx30GbsPSWex3OIJpv5PXmIpvnrYZhf+ld12NYMZVoXkoLXpWTjRBqMpQyTnKkz5KMMknknmoolB10
0JGhyTLHiJLppqMY02kznDTjqTGgivppqKOCeoyRmypJjGh5jv8YjUsDLFPkMjWeVFOWLHYYHoZvrraii1qGuONBy5nI1kxk4HQV
HDaqUeO0lNAXpX0raQboSvCRAQcZXipg3WPLoFdkPcxsBE2n7I7a7rvuxjvKppqyWm+RrRZZYqwGqcGrDpowekyonAzcTCfGdKIq
kgYjrHDB7SIpb8LlvlQlOZS0RAYz9HasKaKIDsWTtJRAK9WMDZoUX2YprTaDjwSMS9Aya9Vm06XoOqpzozzvzOSeNESjG9DLEGkw
qvhqmiqqp5pKKtJEJryw0aE6rKqm5EW6DEdck1MJa5v47Gm+PTc66sPHdCrxvMyMR8AaVZqlgBoIq61q2T7jrSQ9ZAn/RYZZ4Z05
UD3OkSFvx/aumrTiHN/bOONkk72Mb/wmpUmWaly9KcEHj60pw1VLXTDDmjqeN+lxHohaNAIO2XDCRGq+dsJn1672kaW6a3CmCYcH
nQ7FCF0ceUV3Qnbd8R7/uJLKdwxvp2pWTs6sa1Fy6qaiHwN652hvb/XoVHOvuSj+tkTlX6zBjvbts7OP99gEby6/0qe6faIC9awH
3cSd5076MaXaVAA1JcCCRe12HSuSmkbjEr8khXoEgNf/EKY5uxUsdGjDoKowGLXcjSI8YjFIhuKwqdPpzYQ+29QxREODNeggUkdR
Q3g0UaqClU5xRxvf5jQXP/rtTlMDTBv8/5pBmhJ5JhqVUIsFO1a1xPVwh5y7oOwOmMNOyK0FLynR+e5yOQXo4HCcO1sNd+izGpYq
SUxUVcVoAMOZgUUBlLhh6NinvjpK7I5Rw6PDDvg0UJHBK7IR3E+yOEdjQNFUuwvVwITISLtNjElny2PUstic9rikHppoS/ZiZ8Oj
QTGN7fOcEJ8HOwFBxzkEaGNSyLAZMqjqkMhjVw6Xhj18dRJ2nvyc0USRyZi8UChNaQslsNepEsZubJxsRgEPxrFPIQmRnARg/EIV
Dd9ZrCibUMsmzDjBDkbRY+DUlPv2iMYoNoNY29lHMWJSPD7O64YWZKbxnvg/qBlPkkLsxBo45P8V16gSKXvaBOcWN1Ba0nFeTWQS
N8d2OP3BpTTE0KbozKg+PkJynJEMpdo8JjcMSc8i7qkjKd+ZS0Zq6omd1OXDxNdE7RHJLa0x0T6UIRpiqCpipEJex5amTPr18ZPw
jGQXxRPCu7BzjLrcoUiHmEIqntSMiuQEM8wim9i0RQ00a4kacAexiiKPoHebHVjtydVj0Iw7a5GZRXo5N0cZTGoi5VS80KaqMTLN
kKbLaNs685I3EmAf5FiOQDW3pLm606uwfNwTxcmbAcDGREVFCm/2t9JOJjSZFIyfRZE6y3oebIUE4hAcZUIJjb6OpM9MaMOi2U2d
Oi4lfl1LPTLip2P/ri+jDC3mU0+nuA3q8Kmh41WcwgPYttAws51jX10TqkNzWlaeOIwaMeI0Ghr8hRKu+aLacrokpV5UfYpTn//O
ODFR8AqV4lkDq8j6VkhmFnGxpCUQE6q+V/U1LBgpxkt0gFLO3lCOThPrb+WKUGUq8qlFZAtgMRc/5XoVtwUWsO6cW8OKtcc5a/mL
mqxnWJcit5OplSfzgEvY92qKra0Bk+9KK0qNevWw38WockPXDP3M8CJyI4BNF+vhm4JXpdz1sGJ9PM4j0as5LcoiDQC7J8NF7EjI
/OryytlTY7rWk6E663Az9JcSbaKD8b1oGsHrVJfm051yzdh2vITWNRjW/3jinOgPR7naeypsovH9LVsHdJEs1m3O+RpVONNIz2Y+
Lbrf2yFbOOScJZOjyWHlmJRXur2uilXQQaV01GglnrtUjAATLWRyP2lRKn4zzvKMZNqA1qdifYWGKQ3vk0MZ6O8yqYSYNlg9QLRF
iTg0YB9OLMOYKkrgAnWUofagMfr0xto4Wm40gLOQK7tIn5JUd6jm3CFxGaot+cdyqGmGsaP5VFyGeHH3WupXeWgMLp1yNDoSF42R
5z9DU3u39mbe4rCNbH/BJkvWpYhKXGIrKL+LmDm8ZTHnncxTJxyHmppViVTj6PLRoLC6FSJUdTjpY/PYvLxSQAJqo4xVroUMN/90
sSwBDGI0Gzxe8gQVBAcUpxzveIrzTDX75OrW/5G4v6HSr8i/o1aHWHhjPbZ0qd3634ebOrlDhDA9UuQWHQD2ay0pabsSGD4e27PY
0qRzqLhEXQI4sChqopshJbhLD+v11nH+XwD7G1HogAVYrmFr2F4c5sStl7Ne9zHa7jWwPVV3IjnexIO3HueMk1GUh0VgVANMYGNA
cMvQAexQj1lgpdty7lXmKnmXmG52uSC2vut1UWRSDGYgkMoZRSRcxwp4UhepDLO5sFoutxlgtz2oS8ToqG+YxzM+jqlGyliWsLgW
Ykgkizc0aOzp23kr+xjZswZ8M+oh8pdgeC6AFbr/2eu48Vh+ct6qLretmWEhF7DZRUXPyKehkfNOxtOAk8a4oDPe9NH/0T0EQgPi
t00Hhmi59kT351t/Vnp8J17icSwOwXtelHT902B4VVI3pEMQ5kmz9jid0ivF4mgWpgwHRX75V1k8BHUkpVKIJAqGpxZvRAlI0UsX
lyljZHCVN2typXXZU2QEMnI1owCWohbEYFkllYMP51MqeIG7xVPTlGOpI0gLkWPKIEdpRGjIdkhdNVIPJnhqBFPwBlj6EBObQGBz
Rnz4IlYjRkwVVVKRwSYuIYMxJBln01uaRU45hXztM2bXA2At4SYZMhpD0RZlaEZcWGC51oTtwkPYBz8a/xd0opUlXgQRq5GIVjZ9
iFRIdxZ8m3RnIRZrovB+L/EymTc9WRJHx+CEjwdJo6aJYzRjJpVjwpJKAJV3bLeAx5Zr/KaJv6c9OTVUqYFOxcBWxGBsvIgwZNR2
t1h5svNmj9IcMEUDUogQlCATmTNHCTdFGEhoovJ5P2ZlQ+Yp5bIZsRUWgEUOV7V+0tZ42ZhrcFdA4LQe+vEScHMUZChFCCheYNd0
PqUzOqQkmzAsSUYAQfESihdi5taLLCdrdVZmk8dTe6RmtIFKqpcQFmJ27qNC+PZ4fXgM3Hg6fQQ+aLNsLzGPjkYOWUNsK7iNxtZ0
d6WBFrQGMKglH8UUL7EMoP/3ef31ZxrUSWEla99IMC84dCJXVeQggZqQcXgzP4wjMWElX404eOElccE4iXjRHuqGO3xnZI0HV9UH
ZCXIeJmSVanTFmRwjn/kRVy4XOHFkcZ0iQeVQ8onWvt1FF3kAmazc3ykhUSWPUEFdSu1QMOiVUdJiN2FjVCjca1Fbmc2aWI0UFMV
Ii7QQAzBVmogS5j1dn7pXS/3ZGZ2bhfYLtznIyB0jiyhEhcoaBQWYty4hUPUkIsEB5LoEsxXkxYhNzoQlmy4kh3JctoXRuwjE+VY
GzKYTdBBDJV2gyc4MeDUMWyIgiJJJLwBLObDEGmpAMUAYYiVTJZIfbO3cTu5WRX/Q4pYhB7nuB4uoYeCt3Tpl4TZCIqTBYLXhHb+
VDCgV0WPmYxmRl/Gd2ya8irethoiYV8EUDzM2XQ4FTE8N1eHg4Z9FElcwh/zmRCroZGeZDRQtm7YWG4kFUZqqDzIwwm5V3XniHUR
9HZPB5k8yUxux6I/h2xD9X6yEVk/QYbN1Zc6aVcpqInWZ5+hsk8gCIYCYZwKoAz5wnPruVKLhINolnKmhmxoQyT+hhbOlxD69Roo
1zHgySlLxzEFJWM7uEQbmjbKIFxgcpYCwVYFh0dmNl9h54QDZoHzoprNsE785G5GkVUEUIW4dovdpZBZCJzJI0f+4iajAYghlEmu
UYZgLLlUgXZw/HhmjQmnz3UrvsMr/1QQ5aMAmsCBQZlqLgqOHeemPwdV4XU3I9oW/1Y3pFmiDGvHXTGGWCfFdStKVmiTSReSYUUx
VE0ZTX/nmM4UdZSmat5JRDBBA98nhKzqJcvQf4IaT2GWh045TU7zlVHpKSxBnSKSEFzScHC5jR86XwFmZHupUo1SMb4BiJuxquRA
Uy3xZZ4VayqqlXPqTKwCi6UzjibiV0ahBlUlqi12r7BaWOBJZJ46CmdFKS9CHufTRTpWUhZIbd/0PHXGdZ7Klu1iXxNJGzGDEFqm
A1dmOsfXgZuYMBz5dknajevJV97BruiqCWI2ptTqeexJsmMDoLPYEkaRduqnWFioiLLaXr+5KbxHGqthF0RaPAzpXPAFlOi2pPGa
arNzU//VFFo16bDAppLJVqseiW+TprKgWGvlegy7RixwAn4DoRaV4Jsw56Ogt1Q756QUlDR6CohtkQDTWBEcogbT1pekd6F2loKs
KKigUo2i4RllNxBwEB7FQGG5hUuyZE5D5LYUdocNpznGIENo5RIIcZ22YqoOyorq6K1oGDX5VLBNhCJrwa7o2Bx8S7mTGrnlp4y0
tDh0ZGEqRpBAYWHFILesgi+QKk2ga4duiXxqcqkKQIpFZbhFmqT8Fkq3NU4WWkh/+4lZBSwEUJHoGBbhWEywqH8kBWU/CbTZ56M7
BBbhQYpVdY4oeXLd9VR1FpKT+pQy9mKIqxrYCRQEWnBYxoj/Pcu1bZeYvhk/7uZqsRKftvKYOseEqRiWkJuVR+o4nWUkf8RmEio9
+iUbdNOPC7VybieWk3dZKypKHpqi9+NL7NsjKzG1TGuBmOvB5npa2YO7yAqDeTsRMgkdXupB7+m0m7WabTdXpYKuQViSnDoQ5VOg
ibN4+dk0tNOh5afALWd+VTNZ1NkvbbEMfaSGOrlD8Vt9dCRiOTe5FCVLMmFJLMS+HbVcO5qDXtWM9gqdLVjEYFKlGQFtuIhGHrzH
qliCezhNQ+VXGJa/ApExrqEMyzmsK8mCtpqHsiunu3g3ZBcTqrdAXypG+SmyfJmPJAtmWlsk97Uf6zsQOUwAOBh3/07TpwA7a7n2
SS7RaKkRfxTxEhvDxdzWxn/7uChIb0wDhSbSHwRBpIZGUk48rT5MxSL2Vgg0bsjokpsahDQqE66Ej8JrfxPWijzHnbMDSsxzyZ0A
Jkl2koW5FpgiQSHpXw18aBLsx6U2mJT8E8vAsEXWmvWWjQp1V1DqY9n6gFaZpuzUkjIcMca3Nkq1XVA8wprmiUP8b6xxJnWXvcnm
mlLmnw1cSzjYUnrExMwkimGIHQQQSPB7V//TkV/Kg7zpLtRpHrIsEWW6FtlZr6f1lj6HcEBHzzYUoKdEi0j8EvXAmbemf4r5Zujc
pNvIOWryfVu0T2tBUVTmjWX0cW3Kl/8KHGwnBc7EIs7uqgDEgLI0ho8DDUrEup4uNZRfsRaZShEI/LPKDLfxu6PQdUM7dKVEqQDl
+a4E8TUuoQz4N3pBC3TjRc2+e4L4dFIay09Fla3a1aOD5kgrZbKll86IaNCwicqcgxYYJs7l4hKdmnIoGrX49zik2l7upMKOpSXa
i3gyUWbvi2w3CpmM9ESPJJuPtSVaTRDr4Rp8nYaSRkWX+1XwlYnJqFPeG7DHUE3dkcUC8RJJycQIBZdwK8Ucd7n4h1uCRylASKIE
4SJrQL1C7JpB/cjOA3btIje3cdYToSa5+UmOFJdxRqvTx5D4Rt65q7MEMZda3MnQu8OdBL7/D6ycQ/tfyvcrc9Ou7ERlAty175mJ
HeyIC8mYM4aqyaJi7Nu+QljPf+ZNTdndbAnQPoStrSaJthkRmJM8fnpsCRi24H01qoO/NEqk9IBAAv3ZGkq8BBR3GOeeGdhfeios
XuO+MS3CoseTQlswmHw9POxwHH4MYNLRA4He1RaVjTJm8YWIizM7/1fe8mdcwy15rzh7Q4Y0Cr6ci2gipEgD6KvTA+Gw0QDZ2gnT
wBfRSFqptCuVxJaW91ugadfZmSzaBhOpLmmqJViBBtaCNRyEKjLhaUkDCxnjsQtO40ReOnMM6FmUQRjiD+Gw9x20qRXe5bRt28yA
nQCMKKJ6SUwk/58TZflmy6xMv5CHXE/8jabDG8hLVQpA5d6kQ3zJtdUbeQE708jjNhZsjrbdFgP9oHuYbn35nFFeOptKIDRKEYZc
oMgsqJU30s64THnMDAH+4eRxdrddoMpWq28d1RnHXLFGxvv9tDZeJLJ4qV/EjqoV2pTNkLm2JM7Lfyx5DJgnWhOeRK4BZkRtuiT2
tcfYjUMmj0FYOBmB3uKLUjulUBkV51roMYbHaDBlEAJCD5X7wlDp2vn4wXg2512XXF1k5rNhPbO0XQNm06ZqzHvJ3lLODAX8HROO
q1rNTFv45P4Ya25dgN5JpGgFHbQVFuYX0uNe4xCbyl57WJuyGn4VJ//P7lBuqtu4yMk8iEKSFGZOw1ntRnHSLr+Ryt3Yl87RG9ED
rc2IdkZWlbsT/mmhINF/h5jjmn9ud+zHEOAUlxFqsVXSRrlbamxN49URa7CcoK9l5yKZWgZqMWi26sPTO3cZPtyADsH/ZXEu8TI6
4LPr/cVdCKyqGK9AeYujAh4EMuGtC0cdpqJQCbyp4tnklSmBXMQBZxF66ggsyrQc7NqMHHOx5kTN4LBsURtgUXTfzuXWXK/8Vqrm
ZO3xtNbPg4x2c/hqELYApotQHGuL+Fx3qM5i95vNMHB8RQCAOOHc1xKUoCqiZPw4t+c53j8IE2fGGXIygRGGK4CiAHabvnX/Duxe
I8uSAEGGgAIFBGgUJGhQgTJyDR2SQWhs1DFOzTpJPMbs4iiJxjppzLjxmCiKzTgd49hpZMmKHjV2TNnxo0qQJD+urMjJmBqCLgjS
WGYSpTGhKWsOnUjyJEemHoUeZRpSo0aTFl9yXDmz6s2LBwf0TKiAxj6HDQ8qILOyqdGNREWShLsSpjGWOIWSZKpJbNiDA8v+BRy4
Ic+By1DWVdrM5UyZzFjePebSLuOUQkvqKEgjwUAFLRDSCMyzYNWcpFcmbovVJsipp0u21XoyLtuOkGfX08GZgI6pM+UmxWm3tNOT
Vxe/9F0RduvSOSmGhJlbrO6EZMtK1xF3Kcys/72Jyj42E6Ty3lJvgl9GYCCBGQgLCoYfWDoNZuAx9mbdibRVmvrJ++dvrrWmok4B
nwgqiIzAIBqoscfecu0/x7Qr6ju1LLwvtbnEU4kkuTQZyCdNluPwO9MSaw6pDVuLy7vKlOvoNNtGWi+hvgiyziE1BqLhOZJesshC
Ck+qqjEOHXNuNcaUCQshFzaLL0qHaKBhN/5cJA6yJIMzjraoMMLoJL0OPHC9KuNYsC8lLzxPMa6wpOwx8FDDz7+QurOzS4lEA6o7
xOqKiirizHMstqHsTCzDDat6LqqKdCDTxoJyHIwzqlJEcaX9TiqtzZM4VOxFk+rhkTMECZAyymUQJP8DPA6FE4pROjuMzKihVKRM
o0oQ4myAGxUgJjTOXFOK2EOB3G4iCRsVUs4HY0RxU9eqhIPFQ7cSlLLFNnRNwMXwXHTLHy8S6sbNFDhXAUrJoWS6oJwD9b9B8fJo
MhO1ZcwqhM71qbBU4VNmPU0svFQqWMFraa0wA4xJ4Y/mS8DAJlENLaG4amvJTglpzfQq125qCkBymwuVrjXp/YgSDUtU7c+mtlox
3+KOTK1Wk+0FThTpPit13U04U2bm8kgysk7mNp1RYY5KTSiBKhn6N7A1OCsm5DrBzTk84wA8Lt6hgpzOPR6/0kGwMgaagU1sV8bv
OaO1pe1K4a6F0UsTnbP/eu64IPOtRTtFlRlmpV/DcqLcqgyLs3WZLGiThUc2reD8BHwtxoZvflS3X4uJOrS+lCH8omPFw9XFkrqW
rM3G0otYc/cUXDChy2eni6MKQR3UVkO1m3vG4WznNEjT/rtwoqMyJVpZLZfv22ahtH2KsYfRzYzHSf9Kz0BNrIbtRFhFSlgmyApO
qa9fBwq287/Mf9u7518afcKPL2X4W5JA/Ik6xDcRjEEauEsR9yDDFmQVD3Ugw0hclOQYwF1tXDMqmkje5iDgLW9OWgPQlthSssNx
5lwDWRc5EkIGJJVMeeAJmfewZCwCPowArYvYgfinvrIgRAfbGp0Bcae8ef2m/2u24Um/MpMuAiyjfwj6WsZKZqzwvcpvyHJTqOZG
QAzGbGF1k0ymaNc80nEvJEXylFAQ1ozgoEwjDLIR2ooIGMSRwVHRk58SwcM8+s2MOzJBI+IQogkaOiR7BFBDVBLjLAXWBYtOfFYK
ObEzBbSHaaARDGYMgpyCoUwoACQjycBnGoe1TVx/Il3G3hY6+bVsXPeK4NbW5hpBtdAkCTzGfEylOMBgRgE6gJvaLgin4VXohN2B
iB4VQgBK9LEhxOCMGqBlLIzdSU8vE5pWJBO+ZpipRrop2xELEhvfDbBr4+kmiVyDp0pqKppMjNVwMhih+JHIlzbLm7xedCuS3AWN
TmRKSD0AIxCxlGdrrWGLj44RuTdpi5Wd2NHECFJMYwZRAZtQJJ7Y1JuS/e1xRhKcMUD0K8RVKQGVgM/hFOCd3RkrWtQEzzONRDnx
bSdOcSnojE6KFNe40iPNQl6EOImxGcVSIb5q/1IIRUOA/hgno4hSUUD9xp0FtiszBmrdGoxJDogURBn12tsFZ7UiQp6ukxvZFXUK
xEdtKiBfYSxJcBjl0nfKcaAz8x7JKFesnHZJoGACHnQqg1dvka6g82NZlmRzzxpVyYh/qQRnlqGntsh0ijEzp047EYcCKQR2fZyP
tJYqQBRu6G5BQirK0OgZr9xIn5E0yP/WOVN6tXClgtQpeFaywkSRzkOFtI/tAuslm7HMT6JyTn5KKFhvjsKWB6LBV0AIGL0UhBi5
g2WKjjbbkbXTimNCV5luOVXO9Cguj81TYy/SPNjOJZ1SESmv+BKfqpr1bW5aGOoqSB6B4g2WqIElF/8L9zj64WunWcOt93K6MY9p
EUMcgVTiTnXYsqTnIBBt5sneWjoIzfdCWkPKz8TWl8uqrxgIaVXSHLSxufgQuI6tSMCahLgDdbiWPPIYFpWWpeU8SH4+RGpwaZYf
nb72prql6BhddZQ3yel2qtxSS00jyUZWr5HRAEw9EOKIEibrsxek2Wv16i1lnEWY2+0j/hSgBlmdrmaFrClTtcbWQ5pkEwT5ihBP
Rdb+hch41aWoTKh7r+P9aUgmpPEyNRndIs1YZj48FnVHeVuDdcqu9DpZMGt0qtOWJRoICSQsfexaZYEXTjApxqQ9U5Bs0tChmxgj
hLY0xm3FpMwFXMsx+In/oK941F8h/ckz60vho3SqpbCy35aPfOgCH+O3WK0cbB12RXJ+b3JrhcxBmbGzYSIOymysEi4ZC7Mg3Zmv
VlGR4IoXmYI4cr19rGoRr4g0M0tRSff906ozOW0nI65UUeKndycaWvNARtPz7bZEt0qzYf8TNp+1Im8HbEHQ1mmlFlwJYW70lYJU
+jq5hpt8v1lS+94keXSpSDQmXSZI0tA99Ci0zdq5324TOll7nkj21Js/MON6N+8eHH4NPTSv4jS2DHOerwHOzY4zEHNRNG9FzBnw
Qc+zta65Xbp/4h4GlyWhBFB61slINK9VN6cRhLcozqLd99AwGuaTt1axVrTf/+hQT6IgRtjMpxs14JtH05SiKOsCJHd29pcBhvmW
8VWwwfsmxmuTltP3bijjpDmwRyJMI82kgAFU3SFiXkZf6ddzAI97XLFdidicxpk+IpMgZBBwDilUUD3XFtaDOkZCKb+vECmAc/Gx
JQ1g67WnG3h5PyTSSR7ryu3wGsdIk3erITtTnpvRzNny9ygkrVDLN2QTfSEGzhJupJV4unkw0oEjzwe1zr25IGr4FPnwjmxdkyzV
bYOL1MeOIAVU/y/ScUGV5RKntFcUoOFyNrZZrW4BMlrJK/w6KvUzMqE7mJXTmMQzBsICizUCDJnTBDOilaYbMccapAfKCcK6EfvD
N/+E2AQwkp/Wmqu0cj+a8qflSI9h2rABKLWQGggdmLEWiiP+2Dfx2TjvsRWuO5qauaBPCqy1ScG2K5ye2zeHmwjuSKgq8QoEIb+/
WA81UInfyRBYISW/4jtKgpyq0i6fSJ/O2RlicLS2sJfo+yYMSzs8WRV8khgQkxLzuSjgOEAM+7h54p4S+hMogpAFyriW8TUTzMB0
gq9Yo6ueax4KEailgxFRGCr36BcRlA4SWimDOqWIspsda4s1oD3qocDOOQifeBftQ6qX6j4CW8SkMgZK8KCCaJ2EmCHcq8G8ajwJ
qyuPCx10EiwKg6faYTZuEh2YqqNM+q8f65pmI0Zy4Tj/5jmGMpCYWEQISuwJ/sAWFsoS5Qs33/oP/OkogaEh3TC4YkxBGUPG8HIT
KEo3FpvAKRQMLwseTFlAmsK543gWomu5E5PHDIE/0Pu43YosFNwpfjspH/wOZpC9UukLEZw116OVTbKvAjw84fvHqXizAnkw9dGL
gUA9pWM0jMotsPsqJBuKp5olBCm5+OgLHRA3FFSdTTyKn1s1dgrCXhySZCMuoavDm+GdjkHD5HEWqDi6kwBD91Aud9SRhdy0ZVq4
IRSxvIq7BIgU9ai78qvCT4rIYGO0o4C2dZsLmeOMA/GM3UiV+SCx6GmrkrnGADurEmrLjOM2VIS4IpMuXZS3/7ZDwa9KR1pZPvwx
lTIRQQ1biLU0MtFBxLpIuZtsCrBskqqMGknThNS7GATaIpIqp3ochY00tw1zTJUktXTEyd2CLnxRwSjqLAg8ITFSIQ0Bt0JCRU3a
vNlotWQLQNMppKqyNbAQQRUjAMdRKgw7EjakLu9htoY5u6cCKgJwMSlhsk3Ak2FjONwpEQzsr52gjhlgmocqyxrUSppyud/rG+RD
vtAsD/raqr8bHW8BnG1RTSGzC+WDOJdMx8gzSYK4vQqUGJVhu6YcN07oNThKs2OorNNTHz2qB6FpN28Bt6Lrj56qyE7oIIWqkn9Z
ya/iKkDcD7uMT4FqGfeRRxwsKf9Gcw44ghwl25JW4o5i46uDCZI16Jdf+YoxZKPTk6JxxMcm6q86UpHTkCW565zsqZKtIyf/GrIw
CqBm2RJSoT8yaZ0Z9MybA5S76bEmbJlcqsMv6kEry8ebOTJDWqXzhK4Ucrn0BI54CZB8dCppVIhZfLFbOkOJIkn8EjSDnIm+UNOU
lJKNTIDsGL74yjj+I9E0u7ONjEPX6cz4+AkdoKfK0UHq7JvvDJ/KHK+jwTKu0iC3gKIP9bHHMyp5ZC3yRKjaK1Tbgw8GIYDfu0MG
7TcNUR4SYYaiLNTOIVQS+sMyElQa40NHFTyEnI5YFCY6i5KVXIzUixlw0rrgI843okz/78vBjemm3Fo8BWVGNftNXRzTfmSGi2ya
OATWv3CoenjWDJWf/dor23Kei5DAc7m2VJm1SsBEY1vGZVIpAbxFiqC28ykIEZxR5cwbQIwwk/oT2gww+xAZ+eG90hQlM50jR6si
hVlUh+smWOKtk9gR81Gv+wQM8yuizhIZvvq+BR3P85S9SAGWqMEMEWmmDYrURLqZLANZbhJV3ZBKsTDZGkzYbaubyLBLtwssHCtG
VtND5TMUuqyfQpQnjUOqWHOOm+WYkXCqAmkc+DA9BYjMhP1X99tQ8lopbSUiUv0X3Qg43NIreLKxOOmhPws1fAKqEYoaEGvJWmHR
nL0Y4DKP/+dMPsX7CJ7qIqTN1KHAQIlNRodct8H7nrDSDIUYNRnFHoSgBLsy1oEULgB5vucDEE14UZPsVvhYFc5wSixjQu6LVG45
KUpIF/oziIGAg7bNthNt2jlNsqwJsiTyHkVr2t/UsYeDzeYRLuEESgyVF0Kszi/5CDGrLDYFDM4gg1XMKpDtKvW01GZ4Q7EhiMwV
jI2yki1xRIjLFp47Q+jzCEnyKDkzCIyNEuRt2AiKS8D9WAEsMZ37V5772PwyOH9Ep5TyOpZVPHPkCImTRIJQ3PvzCUUdFy9EDudd
xZxFYJNgHVH1iTL4F6dCC7RswwZiTYb9ycc4uwHVo3WdQ1L7v/9fspummJujwCT6utGWA1QstEMt61QdKq4G7FAIeUkTMRLL5Qvd
MN77IzXfm9Q3ermsLLwgi0MiOtRSFZhDil1r0bL+cynj+LA4JDsXWE5E5ch3vUPMDDuaPLiJ3csy489iybMQbVjqgo7bQZpJZdWd
lFxR0LDkTBfqdYinxSCANVLQHGBaaRtBbAYmY5oplg8EcU4K6aEQJjSRcJBCSgni5QuKK2LPvCUU86b3/KX5dU0dfVAfE54aU1rN
SrRMFNoT9LPSiTeIa66Zq5IcLoupOYjF8pTsDT7QSxhQRRaRQpyIcVLBWJV+qYdisxUGFFFA2z2+mws0wk6FIt/yTZD/neS2N03Y
totljxzbYqTM7pRNRPrPLArlw1yN7lugtFKSM66TxIpDe6Pa+Ii7gtAESe7Tjl3hEvGlLJE6XsFTwZCygiAAEia++gWX6OzUZqA5
k+wL9bEh2s0K0V3AK8mgcHXPTIG3u0VGhRudttKbuuBEK4qetnMJQl2x3owPCySe1z2g/VLY83Wo7EwVMXOjC30MvCsNxyNRAPwj
sPggtBDoKiEDnwyXl+Yhg7LQO/zYES1ScT2UpEtD1jW6Ey6+jz7jOZ2JNyvmsZw9OHaI6SgDFp0oDDJY+ZHfxmoiwfTVUISPduHI
ZVLWw5tWQn4jm9Ro/Tk/9eGMG7IohXGsIIkuPEyNa4PxQIRWYvSEvoVdTeWjUqfYuCr70sLEj7We/zqOjg/kvclLVVnTFB6rVTZj
iLvTbRJ6YM7pqFoUZg5WtJvjUzvu4Kdi/rKS7ZxzSWkvOgoiq5P9gDc15GKEsYtQ9tSJ7jjEiybuK7aPI+N4jYlO0LAvGwiGsjlF
xYoCxGRXFcC8U803DMuE0NcpIUGjTmNU7D02NM0L0xnpRYix5OBUmQ6WZOnvYcpFfTyGW2is1o/ZJJ7PQxk1019sljFCyUC71ULW
bQZm0ARhgsKoZa+I2MP13uTaJjEQBdIXqh6pbojoJYBo+JgMQctetFGI1ZADF4taS4hblhLkzdnPXbZ4OlO/82bP5qude81rlKs7
pB28bW0h29K91RjBHP+12VMA4q7eAyGGZ72WbdGvB6yxrTzcsIkYqero03UBH25PUf5HIyxrlkHsmfNjKkYLgN1B5cEKAttCH9tn
97QVWzw0JQvNS4q19h1lnrrU+q2KylYw/4YPbUW1gPsnI1tpU4JblBEO6aAOF2jkskDpTE1ARBLhWN4+O0kokn1RBRcMxqZoDxVS
mYAseZzhvAnEMpZt4pOWLTFXFjdXaGOrfz3gqghu6zldRGeS4TZzEoetnBJItSMJ3GQaDS8LsUYL1Q7OJlbuT4aQqUgwOGuSfDW5
80NWwEIO1wYOJjQxXfLstQMODTSUGzRRnwUs7supHZttpWguiqMOIkfUKsn/NOhhM2PzzwbForjFEwkktSixJQIYmDJtPNjCWgMC
jjJDmM0F6M3Itj6qwg1VRchFoCy7HE6l3MnAEs7G6MCOW+8hEfUEvX/vYmP4GfOxd8pDdHKYD0Z91+4VJ/xa742fCr8cvekIVqqJ
9J/bapC01WduBtNrD+6uuc5BXoPspaW9i4viQqMV0v/LR9DLj/GRHg8nNIFNqv4ETS2F5g8/5wEFqREkCGNH5GUKHDUryE8J9SaJ
D5GTGCSJJ8lyXR/iQNfwxMmDQReY+MCQGPQr2jJDoEymbz0ps4D6pESRyLpu1f9E4nWK8/ABbX53dNea+l+VElfM1+7EIsfyYUlO
/9GPOHqwcC6AMRXxgrfJkev/UM1mn5dgQgjlCotjTpX1SOkLQe+D8Y7hTG+sTeD06x4uN+ByCqUAnSlnD6UtDdrM5PX+1fY2Rwhi
iJtjvWoidMswgRV6h0UShI+wmgFahTgcDI+LcqzyxohLK6yemFkd+O5/US5l6uVkBUBhlPZozV936x2/fs+X5PqMEXZ2jlKkEk9k
5SmPePgh0g0bFwwmoZbQZU1APbRCanrGmA6LVXf4AAg1BBQoICPqGKdmCRceG2WsE7NOxkYhVKjw2EGGBx1S1IjxGESIygYSoKGA
JMGTOsixbOny5cuUZT46DHnRYsiPISNGfNizpsWENv+FRjzGs5PRkMYO7uwU9Cmnjxsf4pzoM+QyqxylVox68CtXhj8f6iwKVmlN
ZsRQDlTgIgGBODDnklTTTCLFqUCjLgya1G9akMw+av1oEi7BtmrmstRBcpNfi4XBUuY6teJkkFofUko5gKCLxApoUGJsmuWyuhmr
NrXK1arkvB97XhzbmiNVoUmN3bWas6/XrhcZXqypBjdEzE4T+kRI2CHOiLFdrw6+t9mmlDRQji59mqVJAjpaV02b3LxvyWgnVkzu
lczJxOHJmB49cFnOjA2bJ+y6/79DljlXXTO+MaODW6DJR9Iy39W3nRrnXXWUgXilF1VgaPl3EHU0TWghboH/HSRdVBZ9RFRTPGmi
Q1kW3uahbOvJeJRZwi2kiVuI2XeSXA46NlqFsAHXm4ToyWbkh8eosWN84jGWWmKDDRlcZe5dVF2LtWnIFFsEfSaag4yNRBAZY5lI
1H4x0uhijWBtNZVTxiyjz4lJ2TYYeiBO6F+eNNAg5F0kHhNWgTO6eFNwmyGaVGdtobSdJmF2dtIyJaJ41W7I1biZgGfxlN1AbyWo
AGPEJKYDbsVduF+KQwWKVGUBaroMQSaNGlpJK4VJ10kRrgqWR0K2uCaFSNa0zCZurhoijEEup56UzdCqwIlX2tjbdX1JCBW2mXIF
0SZthXbYQJWEGS5ByliLYqa8/7HpIXl4HaMoJ2ul5Cipc4WbAA1ksIfmQesWuFe1ffEnsG+UoPRlW3B5t6tLtA5k15SEBnfpmb1d
9BG0ehKjxoazuVdwxq5aGl0nZAwkqJZmXrpcwSR7ZKwxS+LraKQOjknAIzRei+JRUMEr8MXSyndSeMXMJdBJmliL5YhNVfua1Hqx
xxQz8MW33Y6hhQIxTCUV1JxSgQr2rFJDCQhclkJNlBp7sMmrbYds8xTyTvEVQ+XQHjEXIIfmVdXsejVNGloCoSW4RpjRiH1cbi2z
KZG1O/088EfV4qvAl8TMZSsBlRINWOQVwltsbpYRIPbm8YH9kjLxQa7hkB8KCqyqNP+LbIwmBGxyVMiput2RX1bnRZXWBKhbdFQ/
YfZmmh5WO+N+GkO0RnzZD5TAw9+dCt3wLS7rGsGrTU6inwOEF17OMZEUzWBFtleh86slCVbdRkHJuXxwbafr61rSljJEx2Jpokpz
ohY36EHvNtkhg3nWZLqoSSd1WOqPKPw0kE1s5WoGRNJtAMUcmEHtGL1bH9JOsold/ehPZeNbwP5GP8tBZ3IGOgaCuDMQxr0kNSah
AdtkaDkOFc6DdwOaqzrRKPuEJyUK6N7roEQANUxQfBTZTH4OCDjlDMZ8PNGaB/sjxOKFzFNDMYap2kIJLNZIaFgS0dn4xMbWLDE+
iCFA+77/Q4bwLMN4DzHG9G4IwrFspU1MS4ni6PMSHA2EDMNjY9te5S7fWHF+aYEP17S3Hd8F0CUJOs61EAg+kr3rkYbqSUUUhseY
eYqUvGFIYLY4iiUlxiA2stOhMHiM6RCPbW/8F7ZoSYAZbG2Hu+pdugYUPazl8lt3Sk6r2BMurn1mIABsySEppipYsYqQgLRaq7pY
rasYQ3sKSFxKEpCvTpJjTGOT1R/Nw5VsCY1EzYCRcIxDEB1ErTy+CWO3aMM35pDEJHYBZtqKgjra7c5EFHzOhCqBr0y6II+nkWKy
cPnNTv1FLHE6UnNqIiRTKWg0J4HJHk+yt3lRxGdDKqKFyOhR/3tCCTGZLCg7W5IYNfjygKf0T+7IiR5JQkRlJ6lHNyennoH9E3rM
RCZJHGmtwkiwfALKIixt9xEcjQpfPHSQFClRp8qIUZLwHCrz7slMKDGMJDARzfKsJU+Fzo2pxKNcXW2ITMWJyz6LyakPFUCxtT0l
OQQ6JTdHSMiP5FABmkhUDZ85Q9tsxjGg08GrFuq8NWmJo8xEKDAPKTb1NQ1iJuUnPj+rKT0VUTZ88Qszsqe40FgUaQRAWf0shjkM
bbQyFHzWRH4kmu2oc4M5JQc9ElOJ+WmsrJiCkzfhqBfCGPaH4pnnwTYjuEea7SD7awsNnINQ4JmvQ3D0G7eYJxDimv/TXLsKD4s8
2lvqVlCkICUhYRMyrrZ8qXvFOFUJnyng1rxwnPkpSahEk8LjkmN/mvhtm85WN8wd5HSJLexPcHUSe2b1ruxBZdRYybT90kAUT3mu
RmkGs6qYbYQU/kjyunpOKJ6GaQSAZtBcfKm7RouciOrWV+DjqPD8lSWMFKxwLNgiK17onuEzJTMkdq/iro/B5NipkntaGUF508AU
clX2CtKTLGKJWIYyTDrbIq8ya+huuNHtQLF1Q4jYzFYp5OSubKaAZZDOqgpcrXbpV5GK6JmvCrgmfELzO6xmLIEd3eWU0osyhY3m
Mzc9Z5HZ6UM8ytOMdGXlWJeZuYQYaSD/XEPVJFf7ZdaYmRlQ1d6YUQZQ3qptK1qG80VizLrSHhNXmoCpS+Usm+kyM8ILmaY6i+kS
BHFuGZmrUqtZ+0pWdfZOMS5meFbIYMSsgSsVtIlQI5dFTbXnyz0WzepEN0YEPuu17KKrURO0nQHwWVjQPh9rxdm36GnEqI7SkbbD
NC08UmibNTKYlRKVxdweJHYFzV49WsK1Et+SjL9psT+nl+OkRKOJhm5SxI+7P0ocnIx+LOLTZPbiUSzD0rUaLIGG1FAYImRaHw8d
Q1z8R9wipeAdPphis4QgkyiO6I4FG0nGEz97TjirObEnRIDpEYiYMz7adpxKoNfTcNYVidEq/5pIUng0AvBrncdNDCWal2+83e3P
7A7q2cZSjyYhzTpST6oHR7kaZJ5EnbhSRt4zAxbsPuuWMSeUVfxd3CYFPEw/dOFCMRz3EDJa7eLkhNEV0ALQ5Mxegq3QcJiqKgE9
k3Tfkij/SqKjRlp5JCaJA8rKDVwcI4VdW9ZTLy8CitEmZnm7MXjXl/wVDWeSIMUYJL5vWMLIIsWDtxHurvllUT2SxIBmnKP1sZv7
eOPrr9hT4aAU+CqZdsvpl+coJ7LmRE0OZPqdJMkaIxgzr5OHeTqZIUb+W8yxtfaVQguR/0lMEzWJJohQUjHVCwnbv5SRzOwHCvUV
nuVZYjjbVJiV/2HZ3bNxjI6F2kLU2Z3pilG5gDJ4kYsYUf1Vj9uAG8HcxTHEB64U30lZmZRpgqol3OhIRFOUBy7dDeUkBCMNoALc
V+V8XaE4z0Ikj9/tCA2mTmwok5/FEU/kHt98GF5szuYl/4bnQEwxkITT2F/o8dbP4F1goIhFZMc5TRR4pEQ0FVIW9d92kV/67UyY
cU1oWBk5zN1JrEF6LKBTZaDcCBTiSaFCbMIA+lrgNF1S3Q6abU41dVukWR4OihOgvdnfsJT45M1soduegU3sEARP3SBdbRz+0JPl
jBAaJdidkQPWiceA/dNUqcmaeYsuAUUlnNNh4MriKRKDud4TGZ6BFQq35Bt+UV6B8B2+kIHKpUg3cZEJ9dWomATJ2RCzLJT/Tda7
GA9C8NfRUMrrkIQLtFThxRzwBdWbIV8zyNYZ7pmphIYjAcgk8VywCdFYwJm/0R3ruF8nyY6tjV+QIMf8SVJ6Qf8atiwRDVhauplN
Ft0dNoqC1lSTJtoSUKROBHFTYYQPsYHaapSUoYXO61iWAkgWcnCLo6kW2UAdM2Be1eERujiWG4Eamc0P7aGOtWwj6MRHg8igGgnL
mHHYFz7Fa1jVCWLPpTVNT6EX2D2EACqIqU2MyFSV7bTbK8aTRJbkR7jTv3GNMryOMGWFsbDUhBmKu5hZerRQX6lB8lRKAzKLYnGT
XI0PfjmcJorNaNghOehDfFSCa7Wk23SXV45bYdibVHSGQz6jqh2Q2Uykq3VJpZ1KuCXg/VVOoJFbJXpal9zUPrzOShJDZCUiDgKf
BiLgpoDPRegZOp1E8tAAnoxXD+7/oGFdzlrmmEKgHnEhDkroIoONXDzJnpZgkT/pnS9exzFozQC2hbPpklFmjFDYGW0qGARZkuxF
5kYJylnNERulxlsMJ0GEHMQoA658YtE0mcEAoF8S3pXgyFDuyHhYycHEIpTJS1Mt0EYlhNaEmVziEV1eWWh8IhZNmx/Zl0LuTi+F
hZ6N1nyEJN3Y2msI4JDdmQv4zG/ViXRyCpsNy+TgoQ6FRwBNS0Fc1REVYQlOYj8CyhZqpIKRnLiJpXjy0olST8eRXSa64CbaYcsR
BGSEZ6DFCTfhjvMtXf1ZC/eZkw5w5VfeqEUgoYzFB2bFGd8gZo5qUcjokga6U45oYydp/xAB1FAYiVoW/djMRE5UKIWCLajvDJt/
FI4lumXbOeHVDJwm8o/ZMZiFihVTgNYVTSQVjg5x8FadYgSzFd2dPdYBCovUKQR2sh8BKMuKvWZnBgftiGQN8oZt1UqCdRKzEcA8
klXK4Q6tFaHKLYRl+Y+bDgTgTWQTzlFVDQveSQSlTVykFsR9jlxRZuqjEcjdAVsP6oDiDGbSDSGIeFvt8V0mlmZiSE6xiMx0JNzv
5ZWXfkRN0V2tdFJn0JayZMvM5RcZbcxqCNNwhUqnzYuZLqmPrWXIRKGQicaX2Ac+dpJd7lCEpWl5bRFDlE04XU653oo6bccyuGtS
kcyPbI6jzP/A6vgTfukHzZCp/XDd1cClqdnWNUGMGfpOyGiJDj4ZAgZVDkbFqvprmZRiYWmXfxibpM1Nx41dE93Hq5qaJjyHcyKR
BwXkIaJNqiEE9MXlE4FlYdKKSWwe61yapVrIAdLcjqHYdHWItegfXyXboXXS3DVSWbnbbgIgvKqVcODeNy1DIe6IGsQQPtlND54Y
j61m1I7JckYqDdznHe4U5QXbsv5St8Yq6JgEaXWVxaamyGAEpbWFl9wLDegADSQAfkhaSJHPtO0m8GUGcMDlrZgmO50KhhnK6CFq
mjqdB2lj37lFKExth1TkxCJMgIJFRNTithJEcWWala2iJizWZ73/pEeQUuSm1wOG6UkQg8TyjdSE7oua5cqEjOY6Lbga7JNGBif8
10DoatJSanhEZRL1mNpoKgpiDo3kUOYNRD3cGuTil7aIkxedDoIYpGj4mtnWAzqtgd21Jk/aXWY9HWvtlpM1w8xOHNmZRJKGHrCN
ScmGaT3kEAU6J63mpc1uFPMNKi+aFOvYJtgckjIQyu8eaJP1r7z8qNigprwYoZzBVms2E+7oyVImBjHhrXbKaGJ0oYf+yqC9rM/p
YAllBOVyR7zVgwmqpjE0lsvxHn38yAg275dKcFrRWo0YiZ10hRRVXcNCDFRtgl9K1+QtHUZOXlCYJ95a0z6u5bSZUsp1/+mLdKK/
yofZnu1AOIKqCCyPoq4lllBPAcXRKM4GD5yvFOZqzN3rrt4KCVkx4CloqmCq+V/2ecRBkNSdzYemfZJmwQzyucz6glpCjCg6tUUB
Gmgwmtl0NJr8JcQasM4GZx7ptl5bVMIFi9AUy6/mCssoHcWuOSs5tFBurYdC6Nn+bYdJNAZBiA4TXq+EdcQdi97FpgmUXCHDFG8+
muYNoYmN7un4brKl2In2bHAcRwVvZhkwuhu1xVnB1GNKNp4HN027GCh50txhLaumrF4SlgQ5aALXRAgRYkpCCNdG7hRLCFkBQmXO
/eRSLCMpAlNjGgPiOiMBg42dFQn+ieNeBv+RQ8GLMWCnC1Ew6oiPyWEuLEaE6L5g5YYGPWTxzjzYMr7UJHEFBadYh5DgwzWxrmhH
CTfFXp2hrWSS0pAD0xyfszSDpiqgFHolD8uQQuhfKFtTTqkBrmSFhuwoXs0Q6LWnVpymSnBWCkpiZIbx5AFmwyGYE2Xwm8roY+DX
E7OuBbujWsKWmE6ckfmah0XQzNKd/7QE9uDRWkLL2XSWoFUvIyfEd43dPUOMnmlmY2LvAf1HK8kQKUaEMXpiQkUiIK8nJF3fm32F
aLkvOmex/jnWO0clRwkeBzoZV1RLagxT9nxJ2bKEdhT1RIwoYQOs7LSE1lACiynwGx4l8ej0JCn/RRqFxwbvE2CJzcb+s34MmCJC
l3BwkUwnjT15FpOl7WfODQ7B7pcoWhY3mPr4jgUyCzPFDG1QLRteSHKZE1azBLrg0Uatx7VtDr5iUx7iX0+BMdEw4xpqqMYeV10c
9BOTTteC3R9jxPa0hT6s5mWQpGj6pfV2yuqodpt2MF3uDDGMr0Gr3d1NZBJT0P7c3DU10RhaRGT7XRI2CaKNxq+BZxcLY6BJWpcq
BDJdYSYGMdi0EB5rRop6qQkHUh/iKmj0rDfBzGa55xYBIqZSjkx3r0nxy3A3WEpwEH+2C9fJtefimIT1NdVxTWmGxjVBFRmcr0Jo
zc0pSKbZjC35RUbL/14Pv+vKPm5FgArsumpO9StyzrGOp2ZrDrPbWQkqHzndLPYo+HQf0rFxMPXNVLIdjih1Q+0fclMbGiFHteSM
piQB8+19yLOUpTLDdDDogkysDGqcRXD/cd1FnJBt8cvEHNek+A4WTc0j3tJqPXZ1GCMRRxiVE5FQE0loSiVEVGqDK3VJZ/EyIAYR
86EyV0yF3VuQdJYoaFIKNWyBrwkZLN5Q0rRLNMrGFoU1Aih1jlLdFoVGxPi9EEQZHNc0DQCOt4rFybUy9rgHVSCthAa9dagvt6VV
cS1P0sR9F/O91Dg5kFRFCdQcy02O4w6h4J5lQMnRxsc965mhI0RqV+7Oxv/uS8RBI0XiUEmExiVW2/gU4I4CMtGhfMA5xEBDW/iL
j1NxhAbIXgszRkiEne2k9cqa8Kyv2pGfheixMyaduUvZJrAsGwpMBhJsdAJ5jTz3HqvE57BPRdTknS1eUyMTaoHxL6+nOEXXrPMW
V7V39rQ10mWd4bbMLYV6r1JlSFiW0iVgP2v6tcjIdzNkCn1cfpq7THfhkpXeh4/xngQLyw1vmO7tXEgZHhVDYx3pbE0fsEfPDAEf
h1Hoe7r8NIlddzIY355EWfDWcUvlt89xqzGSGrgZJelogKIV6bXmQfRrOpamfdb4qm9Qeh/0u65KS+6vVcTOiyoOh3/zkl/apUX/
A0ygiw4U6wZmL8fqCbVR4ZqsBTd39nEdEjHY/T7Kz8FGly92wogWgzJ1uxvSaVLIosbsz/5tzU0ON5TQwOn6bAMpL8M5EzTpNLMy
Ir2bBpabk+QrCI2ZtEowN3YJc8d2vP1FhBm+4FIyvBYmxhoo462RfeBd2EvGFgEocAjJMknqJdGkNkAQoEFAgQIXBQnoILeQYUOH
DyE6VEaQwKZmnJp1MjbKmEaOnZh5PCZq5DGOxzB2DLnR2MWSJEli5PSyGMKCBxUMVBiRnDIaA3MSBKpgQM6iBKJB3ETQBbOSJWWW
ZAnSo8pOUF2GxDqTJMhjVJlpUkCwoE0FZXimbbhM7SiZlS4/NmP5EqVLmHVZcrwo1yvGujN16NgrU+5Hp11DOk3pNabdkhxH0eWo
hmwCswIVqFG7OW1NopsIg8UL2WrL0FpXRiZ5kmRirxonGrVJkMxmTTrI5qYxtqxmiGrG6jDslTTVvI9NN5NKGvLfwXCNLQV6EDMB
35x50iiqwDlYqsg77S3p/a9Jq65TqyGDsWr4mc1TNr8b9/3G8qg97qY4Vr+CBASKwU7AhpQxSAFNMipNNfPc42uwrTJayTj4EoRs
Gd4wxGknterh77/+DCSgnog0YWqv9t6iyjGrUNJqNRbvmnAqYgiybCzLCP+6bkCHdOAtJMeiuoslx0ii60euGEytpE02KelIJ1G8
qrGq9JrprSD/uvAmhD7kbccdPQNwSCwJq4++BxubykXnCCOGv+pu2lCtEjFkCiFNeKqEt7tidG1Mul4sTjG+DjvGvviiezOom3T8
ciE19NNkONPMi6zMShUsDc2vrlxGmQfzOs+95UQNtbjnMKJTKKF4S8hRAS8kSJO7/CLpOFPnOmk48VqT8hj8jGGrLP54k1OtHvWj
AceCluGJEoSkGtVPClVEEstA6/IuNbFyM6u2VxmqiSBKjnuRrmqRk6+7lhgDFCV3v+vVtb2sgs7X4aSibLft9tutUXAjcnOsJtn/
UzKmrQwV1LtSi8Oqk9gI+M8/zP7lqcNuMSMmrTgwO/JBvTRamLWE6xoPvfYcxMgRhJLlreIdoyFLB2CHG3SmS+dCbNTGkIzXZF9P
zlTFjnCdFK4e68yp1QABVivWAy+dct6Cr4r6GFST1BYkLYEaioBvsSMGt0UVoEStTY6SsUEkU6M00FCNU/tkZjYJsaCiBlqj6YVw
+0nGnxODMl2VIAyNUoPLayxbp85EcmTC5QWpGIp2wwnvgvTZOy1hXUBwJvuEfJw5kzn9rkzlyGxm8qR5Axu7TcggiwY85yTLT8fU
hW5SSyMkjl3GM6WsoARwIsvsvWMfa5nnyNOZSMCr/6qQTax/nArI57X22GequDWLuuA033y7TTp6rl724L7ywfgSFjXlZWw0UH69
v4SDDDWa3SzM8MonrDQZ6WMeUsmNSgkTi/xWlRnNdU8T9dGamozjEsNFpm0OO1Rh3He68+Albm0TmUrIEJRllWVc4ePJ0yQFweuh
zzyb4pNkMnUStgiEX60zIU9KRBB6TFBU0/pVu164OxhBpRL7gRMNKqG5MKlhMRIaFARB10Ik8ax3NgPP+gh1r/6dqHo7e9cLbQKU
GdhJAUy7oUSoA5rCJI6NtJLeTE6kpvIQSWBlcYFlBvKyG9aNWQ/SHrvs48QoLSddPywURjYhMRK6THMx2/8NGZ7zM+hgCoBscljg
DuZGq/EMKz2s1A8dJCzWRSwnSTmjQ1DomHkFUiSadGHpovgi93COda475UJKNJAtWgtrQSyf0Kw4JdDFrnISi1T4xkYAkZxuinRp
orR89szBjUySW7Se1a71F7F0LYE6uaVDBEYATewSlo0rnCGhab4IYqRAl9GPHk24ugEsz0p/w1IAX/g2tV2ve6uSWBI1F0LlfWyD
6KNggwpYz16OZnEEVBLVtoiidTVjbIrsVtm+uRahbGJXm2xfC5cpwU1RrW31oFyrCmLLWx5QAcI8KBdL90SflQukmDTkGsiCoTuF
T1XjbFetYtrB6DXTVrwDIgf/19SrP90FPSXDFs+AUp0ZEIsA+ctoTyiCoJyVDKLqYmaEOsk7x6kkQ3gcCzzDt5SxFGOhHv2qmYi2
wp+JZQDcNBD99iZKNdTMkm0TIsjSN6Evukt3lnRpYLcqigMWbzY5uSpDUIhNeo2qoz8UnCxd88O8LANvRkzpY8nRPWX0bE2X+mNk
4IYr3y0II/q6jH8Aqjn9NIVBWGtOFAHrxbhslWrjgSK1nCJYBw6qlcdQa526pgBjfbMmAyGfz15JxeGGzIOT9auhzMIbbz52E/pZ
Bv+qNNmUsfE0vkOoIFmiJzt1SZwmHJsCdvcuhS0VYS+sFr1G5ktnmq9K0yQcJ8Z2/xDileUotHtsbA5Us4RiMCvZ++B9/SITUbYA
gQpQ6SnppABleDCp7hqTS44jWUSB5FkZapWB9wacsaRQYTZb6AtXck5OuhGp6/xZ0SBXF/klDSdmvKqWCFAJ9pnWdL0jZBR95xgI
6cNuLFMuaMNUDJQp9GaFNeh+m8hQnQkPpWWJ7d6kUxH7fDgq9prkmT7oGHTJVV393W1xq0Un9u4GIcv9Jlss4znBZVO61wvs8w4m
St5QeCwXPmOGieExq/XXlTNp132Fehjg4EQ/YzyQCRFMBk7ik1SMxvIn6wPWfA6OvKjqnTH2W2LGSow2oF2IPihSiSiWV82Bg+tH
ZIyVl/+EsVt2PmNNDsJRVBd1g9I9V6EUDEErlpggNazIDbXrLqkN0o3WzC/CImyvoD2Pzzd+bzcxg+Kr1u0gmhAssZFT2D4JF7/Z
ChUB6qqbzz72gOK8Zm3HWlMsZxFx6gphdbwm7qbpACjLwLchG4o13fXPXeVV9iXPexX8EDc5G+nyRavqanIAea/D1myyi5M+TIln
uPsdFsZ8fcOnbWJt/L1g6nhpJpJl8otczqmBNnFDFRNgw32pEBQbvO0JNtRUT3EeL3tYXCx1gk4D0Q7ZXKVxBMu6o0buNsMPyxwF
lWS9ODmIoW8osM5xuJweb97VSEqlHN/lIgIt3lA0ZsJNUEf/DUQbHWIg3ka958ruSd+qstVeXMMI7+Y1+rrGydGhgzCRi1EqlEug
48cPVp2sGC9Iyk24cq7+XKzMQ+zPnVkSxwDH8Dca2BnJQgYKNbOrNPXZxNX+cwDSdLdGtWmZ6Dyd7Mbd1U+jhFwz2+d7br1IQMVS
y44IWjcRhBij9p8X+S6yKxpUTWXIrvwE3jT9EKBQSr12Bzs5e6W3e026imHfR2aYklzM9HYaCFIQTw6yCHnk6Hrh8O/b6d+ORAd3
5I0iaQC0gMwiYk45qGiX1Mm6hqrWOoH0kiU3cu6GeuQgvqvDXs+8bIokcIemrAa9Akid/MRnFCuBckpm4q9AxoVQ/75CjrrigsTP
A1Xo5XhtUQLwsWikIMYJtcSK6PDlsB5I/Uij7RqLWHiPp4yH2HyFViAkZLRNt4wBjqiPd4ouaBakcR5k5+yIKIICrc4oGnBC9XCs
83JMcG4vjljrQfqGdWrwx+5E6zjvfK5neuoD365HePxvIHAkAuOpdYzE59IOx2SkwTZFwSAv5gZolb4jJKhqKKLKqnovqzgsCRtu
3+DM1rLod3oE3BzLBiniufxqzIzmqIpEZ0JFpJZJoFAKKLKvaUjIDOtliNQvCusL1ULO02DOAA9uJH6JKpYhqoZnP8Yi/siByQqi
EqhlLrKtuKbt4frMOwiuTpZFAA1EDf8o678iaek+UTTWjcFaSyBYZT/00L2oIxpgaGEEkeFuh2pOLbxkzpe8iKkgbC+Ybaowg1gw
75bYYvFKkZng5UlgLKSW6eOMYWwsjZsE0Hjmphq1br6CcLqmxSpsjlgMYhUB5lnGxaEkUfik52A4opnsrzxG5iFfsCr+rbH8jwCO
B/E6pCAoofgYMPwK5U9cMAwviPBeSwBxZA00ohBDrwcX5+BGBxtBoiSHYnxOiY9cQA0s0ZLM6QC1cfqCbiej4q9mbgVBZiWigc4o
LWkc8RFXDIbchdOmTfL+rvM4we2oyjKgTCiU0lewxKZ8DkIyhYOwphOQhZTerghNiDdcYOH/ao9h3K1hygxRjm1UbOvl7CJBJgRi
hhAnhHEhJGYNABPd3HGkqshX/oYkSG94qAq02ulrRs0VHWwqGu4yt4oZSpJ1wtG9eIOeytESgSr4pI2NRCN10DEQq4wwRqIMBKJL
Omve4u9CBiKFWm666kvvpjIM6yKEdK9VDvKsGBBuxGusToWceiVdkMbrtkQvwwdSVqxhkJDkLIigUOvc9svBHtJSpAgXMUKg5C03
KPKqKMIt+kIgFyRNrnInX0N9GK5x1GuUuvKULmTxuughC9OHcIw8S4PggJFfQmGlyKLuGvLFSk3/ztDuyuPKTkcqseZ0TqpbLC1A
pVFWglL2AgfXIpbRO57nSkRPxZJlwBJA+RCCGm/HZ0Ym26STl8avTMYG+Qr/gjvDx0BcwPzi5nriMTe5Dab8R3QMhX8ezPO6R7u6
ZQ3jrx4SIG9SBgSR09SeSSu6jRMowRfrRERVjg8Rx6OoSO0OjsayrSRraCDIdG/6Y2rWU1vo8lJuC7ymR+0W8ppQJ46qQl8U6Tct
7DEXIvWwhPYQcDF+RxLx9IJYYg1uxGsGQhlscEbv70+T8wxJ6ymJbb7uclHIIk6bxj01Qe/SxIuGrtEKcyZ3MFVDxUkgg87IgrFO
z1BnVIqyBvBcMD1Jhsx8axTqjWXIwsfwkTd00ujcBSouCUoi8caYI3la5UrHAkg1RzrKBj87Ako28NoSa+/Y0UhjCXrqQljgRH4O
/4IeDFX+zmragsmyspFF7esTSQNba8QmrDXzmII+928fBdItGw+dgEcHB1IiUYpUAYYe6kx7QrBtFNQ8m0MrODDyEIUBP3GbTEwL
o85QMYMM0MlPoQeDIu6VCoZTSqIYFOkgvmc1TwnWCoLxTsewPilTHc3FDHE1pNUblSbjvglDMnSN0rRUemdFFKy0NLUJowhxuiJ5
IoZSX3ZdycFlhDJeUUZoZu2VqIIxMWQ4H+tpPFZnCjMOT8M8XInajiETG6soEBZggMO58EuFsunT3Agjm4rWCnAXmQFpMMTZjDX+
Uq/xDmlTRjO8rIeVOEw4iaI/ioJlz6gXCUIHKFMuCf/oEiMDzZIDSgYS4MyKZ2+piFjSgRQQnV6JlZYwR+kQ6Rbkg1ZHO3eMBkzJ
UCVmPWrsiv5EsNbkwcBjJWYojHgjJe8s9Saxz+RWKkLnt/JibPhlPyz1mz5TBz72zSpImNamD5+IYEuzniLRr/rJnTYWV8/KzQqq
CdWH9d7wJBBlFFYSTiL0sZhMJ2TzKksH4mRzJF1xN7SyILzrquyXBvBzNkXweYdo8ngLF8+E76SwR9nrbjAKajt2XD81BoFrARdV
epgBGkB0VbiwaUSJDBw1SZFkLrKntl5OMnADGPdjbQGGaQ1OliDIXSJRNGVTQ9dxMt/FNXbWnZQHatmVAHT/UpW0CLzurWx59R01
gqpmI4ObhiKcV6kW7qgWsn1ecICMIfcURTtQGFzggCxmZV2k8EaHNseo8GoXrYWb0RhukIwITSh0OGppQzaNTWcMNP0uaFk1kBOa
jQaB8860iyPhhnbVJojt092agST6g5QwpGv1w3nXJAMDk40+Qo7285diUWSBZwlTAw50CrmeTIdlJjchDooiaMFY75wwQmmSpnu/
aTtooH0KmCEZztRsNOlGAkSoI38zynGVS7KUUWALCLNUNDmrcfas9y76TzZYt710OCggachw0XBF7w9Jo9NAiahc4EOVBuzCZ0CV
C7P6Cj0TUSrD6vzojNBuRCCu/xhc3FPKLJaAQLHUosmHazTEFip9RymHdRgoIDdrshFpmxLUiHaL9HYs6BGVb4llYOgYjSr66NV0
0RBj6EwBzvlVXEsTTlSOsu2DJzTHUkeY5XkxtC3DFsXScuIeXa0sPPatKCkSK05VF4rRcON7tvaxTOrykNBLXW4/f/ng+HRbRaEZ
yAilCCBz2LDQiGafNlTRTLHD6jiLXEJuoMc17LCesVnjEjl0mnSFIsmXu490IKwkBOpKa7UgQAtHVpn6jmpTjXeSPabYvo0GOTej
ZGcUmkph/tWmHyhBsKdQnkQkU6QlgLGWDcIy+vYx9eNr9TkuwTZo34yShYovkmlL+P8DtGQHHslTZs3voImPWZNlO7RQAVxNhZ30
nzNzTxnmfFo1goMrNcSlP0aIICL6m8ZGB6DwPA30jcsu8AIHLISnKBECtGwCWksDc6NLgISJoMJDFIYiRGz5qnYuB3P0MmGK2Eiu
8ejrUf8qLzKsroyYBkja1ZJlA3e1TxVzq1+IJ+fWPaHOsx+rOh7yMstPkofIXatGBWeCVt3J1VbHwh4Z5PAWIOE275wyJkdGRkJ1
lJD4sR4aTZFjTf0ZdIlKn+01jOZpvXkDfbQtqPbZWugC8NBDPwYMgV77VYyJlShZDgFNn53KOYIV4sRLJfql65hljRcikZO0v2Hz
eb61uLH/ZDMLkizGmiy6gxnMOhs/DMmmZrwg4xs7O2I0riTJB18YFiD5+n+fSbXCuDncxFaRa0RkvEcluSpF7E8zSPYsEyU0wVZz
I1/Dx34HwGpUlHSXuuRGzHhBwhfvlwA0rrld2efA2aZHo7TvE4qfhIeoYufAusKmGvGSpcWmbGaDHGA92rSZo240EaIpnCB+LkKA
SfqET1efWBQwpk5e96qAzcIwdG4BMSHRCbjP6ycHPDFSszJWBQ5knG9wxGYHJ3RfyMrt9HzN+KL4JT41J5mGBGUa/dqGDq0vmyRu
Ao++EfFsQgGpZ5cfimI9z5X3+09iRrnJiCC4vMupA7jc2Yd4/4kHS1tKYoUeW4ZxoQ0hTjxyyFj/QvgHMWZZ8LzJyYIYKOlwgvl2
F9rGlrJk+REmlm9YTm43an0hSlLXNm2Ebxw3P6lXveIzGwuZM6qYVphZG97UIqriCPYjOeLQd/bpEK97NmEZzZIb6aIJobU00YVC
3FM7dJaUEh3x3stXvS/e4xzHW8w5hCW5ehfBEaIFDwp6RljBghXrNs+Ot6TwEO9pVvnDHsymb7FojhOKMJcjjNh7FIDdDdU9rTMw
FRrSSSYKaZgTQKTZdgOvvmn7lkG8ny9FEpImw7WUa6RbDiL+FmXGGsqSWVok9u/KnK9SzLXp6VGsE54c+iaoed6Aef99XFe17DiC
AAjC0r6n5gGGVvVhKgJrcMA5Kulc15heO4ci/kxVIU+z5eJqaIyaGwUXt1iie0RaSrsb8UrS4Jg58CwaaNOtS//cfJuhkE/uwF8F
bdlqtcTeDwdxi5TVMk9OgYOx5BES4Nn5B8mMcq0Re+8jPBafWG3Cd2Wcy+oBjkXj5cNeR3nV9kCCVZR7kzMqDUcrnNkRKpsp/Q13
/StMfqo02nLfyAGiE7Njo4yNOibqmMGFnZpxImhMYCeExyQ6bCaR2cSEHI9xokFAgciQCRSAFFmMnMqVLFu6fAkz5ks1I0FdfJiQ
IqeLC286hIhx4M6hQSE2fKhx4SgdIUnhhlRAQIfMqVRVkhlZTOfPhEMhHkyo1GPQTgw1miVbUCJFBQlCtjD5lEDVuXNpQFWQMKPe
ng/TJlWb0C9apWUHH/2Z1uGmpy5ENibZmK7kyS3JhHSxrGvBxDjFGr5JMXDEwoS/VhzI1C4B1XZpUJ7MVKQyjRfVni2clGJY0Rsh
asUZdyTUGQpeG1dJE6omiYWP3RZsrHbujh1/3sYYHSnzi5RKBg8u9bj4mMkVEJtoVixRwxqdC1Sv9T3zzWVDX217967r8VOvQs3a
VVKbiRXRQJyBVuBgBSFmUGjHQKWaAo2N/8TfZIuJRMZZv3nlHm1dUUcRRQp6dtZzDBmT2nAinQSSJhW+qNIaqxGwiYARbeWeTx9y
WFpPZGEHWl6WrahiSDDCdFVIynCFI5M7zVcaRX4ZZWKIy+RnF35GHlkVi8YAtZODne2UIE+HKUjYXoJRtFUzwd0V3DJc8qfJXZoM
COKOv1mHZnxpMigQTSc1ReGcLKmxmgKbXCTll/QJltRPaJrWIFGdcURAAoPqJ5ehMkW4jJMOdgRdkwpBKamBX+bVnpPYadIUiyHN
ENJ+nhpHyVObnFiQV7gdddqIT/76q0YyjmSXhCLditxdxOwYIHpqDTvKQKKCyOBBQykDJ1zCMf8LU3IEaKLeWO651yBvHjpYW5ol
6mXjKDRNKBycaoD72iYjqTHYWK0es+FZ7T6aLonbIVXnUwPkN4CctyIK1aKd4enngjvKJxR2pXblpkls1Zoovi4VcxcZrJIFZcrr
kpogXyJGql0zDS7E8Em0iqSoyJPBKqEaqqK6nod8pkdmQwVqZfBPsLKIrAIOe1oJzprMDJ1p67kLL3vvQboMAbFCVZJJOreUKAFm
WusgwTj+u6NojzZqWEL1IPvmjPWMTVcxT5FBJWDwgXXgmBon+G9HH14o3IQEEMPs0jTudXa/V1s9cHY4PqnR3CflVyveK+lw1zI+
oQUf26YWnOqP2hH/nhFFiCYLYewKhOc5VaA0ZXKlzVS8KoF+Exz4npgTA6HCoTNLU0jECPiyXj8VnbHu1Iol5kXKMObYSLTjTcld
lPjN5N8crqwQpKRTm51RW/l3GedQ3Vs7VcUkq4NvhvNkGsXR6g70MYwWfSWwyW5czEIcv6o3vd00akpkMl/BCiQc/Ognfl0ziQ7c
NSrdqOww0XrXunrFJLIwDGdvgdDT4heTK9XqUgSqWssOQr76nOxyN6ng5tq3icZBqBLxahloXNgb+uwFWNTRG2tiFxUUqoYA9Fng
0fSiFdKY5j3Q0pH1sKcshqGQKvoglOVId7IoZq1gbmNdlHZCj3rRiwDw//OU3jDEKFYFy0bzKVdnLIK51mHOOxHsXPxiQwBleIYh
TrSYuVimphY2R0BlIBSc7NIYMmyRKiCDlB2lQ0RFHsx3z/GIVkBGNwVQQochoQRXKNLJhKTHkC5j0tGmxxUbghIk28ObJpJlsgyq
8jOC68zRAKdAi7COGaAzSVuSFRc4THIqESLX/fZnquYxL48MAlIDjSGShYEkbGxkVtfssgZfQSkt0vPTYaTYu84woyl8nCAFCRUV
NWyCHuJLCI7IyKE8oquJHVFL1x4zEgmG5G7LjElsFKCJnMByjuNkiMxMs8tWSqkYmtgExNwXEkne6npQoYTA+kI1rOkyWJgUmv9E
VMiauNQSb9kTTkh0QAY1aKIYyhgN8G70QnQuYxnE2ERF1aAGMpBBBzrIUuJWFLKCxqR7UHGBGioBVEpooqJT9SlViwEKii5jpsrY
aVeX8dWvzrSqQBVqUbtVNlnZZaVHUuHs4ECJNQRVqHMdKlHvStTWwKU1iYKkx1xw15jKdBPEWAZxCKUpP8aPGKCbkUtd6rXZEXWo
dI0qUNVAibLeNZvw7BZUCEUvOJWkjUp9SRYx6rXUYhStrF1t3bz1lHq9FFxxcWS9PpvayHZ2t6zlLaEi9BRboZCxs/QsZB3p25G0
TwELe2x+YEurI9KotDEhXksBijMipbVeTPMrSCL/pDhkSTClCijharzpLWRiNFnsFQ4NwvbeFcEXP7LyWGx/685l8uyRdNMUzgYF
uwgdkXNH9Kuy7itbsVE3hQcNbvuQe1TmCmdhdUOmdAMakuY2rTjMai4NaABT4+72mHB6zIwenGD+JkpLcGFr/Eg2APWusXiyBQko
43LhuzT3txSOUHlXtGCq1GMZPo7thBqTgAlJ8DERkuCO69VcJiNxU53Cm4O/i7ObqaadOIstpxSg5Vo9VswSpG7ylMXH8CIxxQdO
8A0lxM2WOtZ7Qa4KFt+SKfdlMbQgMzCSj3xbPdcKP7Vr88cMjT0AY++1Or4LcdpcNo2WdsaJqpmE95y9/+Ay1y4xFmClvxackiju
hHVGkoQrvGb3mTfFxtPPwgD9Y86RGl9P4WNwyMupkmyzVtosbtmwBJ4FQ2DOYPv1l/W8Kfdh+WsSri/2Sk2VKyUL1nmeJSjXXF9n
szrTdlGG52jMV+feRaDATjCLkbhkx7gYhcUAnagHmONOy9vTOc6whUco5tlCeyrEgHB+FMfHHTe51d+xL2tI7GUCpARv6kXuo11w
zO0am9femtAN+drOdW+xHpvwj2dvjeosgjy7yB1whk3ion1PZRlM0XKnQd1w91o7ved+MoRezZYcMhzdb04457zz5uc2Ws+AlnSp
66EJNZz15YeVeYpfzhYpz56K0aNWeVXaXeQsfUevBGi6nm0+cxDrYLBdjZ9/et20tMYWxGwPLGUpW1S2H3zOKrV6S7Ya17yKvOJy
TnCtAhvViva07LO2e0w0UYmzgtq+lw57VNj+9qgSNhoL9mkxLA9Wmi6jHtAgR+H5s9VNZHaodjG64WPSVWVY1aeUgCslCKsJni5j
H/o4Pb7A6lPWt56ihN0pWPfBedsLnxwBAQA7
"""
img = tk.PhotoImage(data=icon)
root.tk.call('wm', 'iconphoto', root._w, img)    

my_font = tkFont.Font(family="Monaco", size=9)
root.geometry("1315x500+400+400")
root.title("Verein Manager")
root.configure(background=background_color)
root.protocol("WM_DELETE_WINDOW", exit_prog)
try:
    VereinManager(root)
    root.mainloop()
except Exception, msg:
    tkMessageBox.showerror("Error","%s occurred " %str(msg))