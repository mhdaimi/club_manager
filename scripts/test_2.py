#!/usr/bin/python
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt
from reportlab.platypus.para import Para

inp_doc = Document("masjid_template.docx")

tab = inp_doc.tables

row = tab[0].rows

cell = row[3].cells
cell[0].text = "Mohammad Daimi"

cell = row[4].cells
cell[0].text = "Winterstr. 19,"

cell = row[5].cells
cell[0].text = "60489 Frankfurt am Main"


para = inp_doc.paragraphs
print para[0].text

para[9].text = "Dear Mohammad Daimi,"

style = inp_doc.styles['Normal']
font = style.font
font.name = 'Times-Roman'
font.size = Pt(10)

para[9].style = inp_doc.styles["Normal"]

print para[9].style

import arabic_reshaper
from bidi.algorithm import get_display
txt=u"اردو میں اپنی مرضی کی پوسٹ بنائیں"
reshaped_txt = arabic_reshaper.reshape(txt)
bidi_text = get_display(reshaped_txt)

para[10].text = bidi_text 

para[10].style = inp_doc.styles["Normal"]

#format = inp_doc.styles["Normal"].paragraph_format 



# print format.space_before
# 
# format.space_before = Pt(12)
# 
# print format.space_before.pt
# format.space_before.pt



# para = inp_doc.add_paragraph("Dear Mohammad Daimi,")
# 
# para1 = inp_doc.add_paragraph("This is to inform you that you are a fool! Ha ha ha ha ha!")


inp_doc.save("trial.docx")
    
